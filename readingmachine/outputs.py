# Import project libraries
from pydoc import doc
from lit_review_machine.prompts import Prompts
from lit_review_machine.utils import call_chat_completion, call_reasoning_model

# Import general libraries
import pandas as pd
import numpy as np
from docx import Document
from typing import Optional, Dict, Any, List, Union
import os
import json
import ast
import os
from typing import List, Optional, Union
from pathlib import Path
import pyarrow as pa
import pyarrow.parquet as pq

import re
import unicodedata


SUMMARY_SAVE_LOCATION = os.path.join(os.getcwd(), "data", "summaries")

class QuestionState:
    """
    Container for managing research pipeline state.

    This class keeps track of:
      1. Insights dataframe - traces insights back to the `question_id`.
      2. Full-text dataframe - links full text to a `paper_id`.

    It provides methods to save and load the entire state object as a pickle,
    and to initialize a state from a CSV containing literature data.
    """

    def __init__(
        self,
        insights: pd.DataFrame,
        full_text: Optional = None, 
        chunks: Optional = None
    ) -> None:
        
        required_insights_cols = ["question_id", "question_text"]
        if not all(col in insights.columns for col in required_insights_cols):
            raise ValueError(
                "insights dataframe requires the following variables to initialize: 'question_id' and 'question_text'."
            )
        self.insights = insights
        
        if full_text is not None:
            required_full_text_cols = ["paper_id", "full_text"]
            if not all(col in full_text.columns for col in required_full_text_cols):
                raise ValueError(
                    "full_text dataframe must include 'paper_id' and 'full_text'."
                )
            self.full_text = full_text
        else:
            self.full_text = pd.DataFrame(columns=["paper_id", "full_text"])
        
        if chunks is not None:
            required_chunks_cols = ["question_id", "paper_id", "chunk_id", "chunk_text"]
            if not all(col in chunks.columns for col in required_chunks_cols):
                raise ValueError(
                    "chunks dataframe must include 'question_id', 'paper_id', 'chunk_id', 'chunk_text'."
                )
            self.chunks = chunks
        else:
            self.chunks = pd.DataFrame(columns=["question_id", "paper_id", "chunk_id", "chunk_text"])

    def enforce_canonical_question_text(self) -> None:
        """
        Ensures that state.insights always uses the canonical question_text for each question_id.
        """
        # Build canonical mapping
        canonical = (
            self.insights[["question_id", "question_text"]]
            .drop_duplicates()
            .dropna(subset=["question_text"])
        )
        # Drop any possibly incorrect question_text
        self.insights = self.insights.drop(columns=["question_text"], errors="ignore")
        # Merge canonical question_text back in
        self.insights = self.insights.merge(canonical, on="question_id", how="left")

    # ---------------------------------------------------------------------- #
    #                            SAVE / EXPORT                              #
    # ---------------------------------------------------------------------- #

    
    def save(self, save_location: str) -> None:
        """
        Save the entire QuestionState object as Parquet files (one per DataFrame attribute).
        Handles list-like columns (`paper_author`, embeddings) using PyArrow array types.
        """

        os.makedirs(save_location, exist_ok=True)

        for key, value in self.__dict__.items():
            if not isinstance(value, pd.DataFrame):
                raise ValueError(f"Attribute {key} must be a pandas DataFrame.")

            df_to_save = value.copy()
            df_to_save = df_to_save.replace(["", "NA", pd.NA, np.nan, "null"], None)
            if "paper_date" in df_to_save.columns:
                df_to_save["paper_date"] = pd.to_numeric(df_to_save["paper_date"], errors="coerce").astype("Int64")
            table = pa.Table.from_pandas(df_to_save)

            # for col in ["paper_author", "full_insight_embedding", "reduced_insight_embedding"]:
            for col in ["full_insight_embedding", "reduced_insight_embedding"]:
                if col not in df_to_save.columns:
                    continue

                idx = table.column_names.index(col)
                
                # # ---- list of strings ----
                # if col == "paper_author":
                #     # First make sure NA values empty lists or lists that have [NA] are all converted to None
                #     df_to_save[col] = df_to_save[col].apply(lambda x: x if isinstance(x, list) and len(x) > 0 and not pd.isna(x[0]) else None)
                #     # each cell is a list[str]; Arrow handles NULLs automatically
                #     arr = pa.array(df_to_save[col].tolist(), type=pa.list_(pa.string()))
                #     table = table.set_column(idx, col, arr)

                # ---- embeddings ----
                #else:
                # First make sure NA values or empty lists are all converted to None
                df_to_save[col] = df_to_save[col].apply(lambda x: x if isinstance(x, list) and len(x) > 0 and not pd.isna(x[0]) else None)
                arr = pa.array(
                    [None if x is None else np.asarray(x, np.float32)
                    for x in df_to_save[col]],
                    type=pa.list_(pa.float32())
                )
                table = table.set_column(idx, col, arr)

            # save parquet
            print(f"Saving {key} to Parquet...")
            pq.write_table(table, os.path.join(save_location, f"{key}.parquet"), compression="zstd")

        # marker file
        with open(os.path.join(save_location, "_done"), "w") as f:
            pass

    def write_to_csv(
        self, 
        save_location: str = os.path.join(os.getcwd(), "data", "csv"), 
        write_insights=True, 
        write_full_text=True, 
        write_chunks=True
    ) -> None:
        os.makedirs(save_location, exist_ok=True)
        self._drop_unnamed_columns()

        if write_insights:
            self.insights.to_csv(os.path.join(save_location, "insights.csv"), index=False)
        if write_full_text:
            self.full_text.to_csv(os.path.join(save_location, "full_text.csv"), index=False)
        if write_chunks:
            self.chunks.to_csv(os.path.join(save_location, "chunks.csv"), index=False)

    # ---------------------------------------------------------------------- #
    #                             LOAD METHODS                              #
    # ---------------------------------------------------------------------- #
    
    @classmethod
    def from_json(cls, filepath: str = os.path.join(os.getcwd(), "data", "parquet", "state"), join_str="-||-|||-||-") -> "QuestionState":
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"No folder found at {filepath}")

        files = os.listdir(filepath)
        files = [file for file in files if Path(file).suffix.lower() == ".json"]
        if "insights.json" not in files:
            raise FileNotFoundError(f"'insights.json' file not found in {filepath}, cannot load QuestionState.")
        state_df_dict = {}
        for file in files:
            full_path = os.path.join(filepath, file)
            df = pd.read_json(full_path, orient="records", lines=True)
            df = df.loc[:, ~df.columns.str.contains("^Unnamed")] #Remove unnamed cols

            for col in ["paper_author", "insight", "chunks", "pages"]:
                if col in df.columns:
                    df[col] = df[col].apply(
                        lambda x: x.split(join_str) if pd.notna(x) and x != "" else []
                    )

            state_df_dict[Path(file).stem] = df

        question_state = cls(
            insights=state_df_dict["insights"],
            full_text=state_df_dict.get("full_text", None),
            chunks=state_df_dict.get("chunks", None)
        )

        return question_state

    @classmethod
    def from_parquet(cls, filepath: str = os.path.join(os.getcwd(), "data", "parquet", "state"), new = True, join_str="-||-|||-||-") -> "QuestionState":
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"No folder found at {filepath}")
         
        files = [file for file in os.listdir(filepath) if Path(file).suffix.lower() == ".parquet"]
        state_df_dict = {}
        if new:
            if "insights.parquet" not in files:
                raise FileNotFoundError(f"'insights.parquet' file not found in {filepath}, cannot load QuestionState.")

            for file in files:
                full_path = os.path.join(filepath, file)
                df = pd.read_parquet(full_path)
                df = df.loc[:, ~df.columns.str.contains("^Unnamed")] #Remove unnamed cols
                state_df_dict[Path(file).stem] = df
                for col in ["paper_author", "insight", "chunks", "pages"]:
                    if col in df.columns:
                        df[col] = df[col].apply(
                            lambda x: x.split(join_str) if pd.notna(x) and x != "" else []
                        )

            question_state = cls(
                insights=state_df_dict["insights"],
                full_text=state_df_dict.get("full_text", None),
                chunks=state_df_dict.get("chunks", None)
            )

            return question_state
       
        else:
            files = os.listdir(filepath)
            state_df_dict = {}
            for file in files:
                full_path = os.path.join(filepath, file)
                df = pd.read_parquet(full_path)
                df = df.loc[:, ~df.columns.str.contains("^Unnamed")] #Remove unnamed cols
                state_df_dict[Path(file).stem] = df

            question_state = cls(
                insights=state_df_dict["insights"],
                full_text=state_df_dict.get("full_text", None),
                chunks=state_df_dict.get("chunks", None)
            )
            
            # Normalize the columns - first get arrays to lists then get 
            question_state.arrays_to_lists(["paper_author", "insight", "chunks", "pages"])
            question_state.normalize_list_columns(["paper_author", "insight", "chunks", "pages"])
            return question_state
    
    @classmethod
    def load(cls, filepath: str) -> "QuestionState":
        """
        Load a QuestionState object from a folder of Parquet files.
        Expects one Parquet per DataFrame attribute.
        """

        if not os.path.exists(filepath):
            raise FileNotFoundError(f"No folder found at {filepath}")

        state_df_dict = {}

        for file in os.listdir(filepath):
            if not file.endswith(".parquet"):
                continue  # skip _done or other files

            full_path = os.path.join(filepath, file)
            table = pq.read_table(full_path)

            # --- FIX 1: remove stray arg ---
            # to_pandas() takes no filepath argument
            df = table.to_pandas()

            # --- optional cleanup ---
            df = df.loc[:, ~df.columns.str.contains("^Unnamed")]

            state_df_dict[Path(file).stem] = df

        # --- FIX 2: handle missing keys safely ---
        return cls(
            insights=state_df_dict.get("insights"),
            full_text=state_df_dict.get("full_text"),
            chunks=state_df_dict.get("chunks"),
        )

    @classmethod
    def from_csv(cls, filepath: str, encoding="utf-8") -> "QuestionState":
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"No folder found at {filepath}")

        df_dict = {}
        for file in os.listdir(filepath):
            # Check the file is a csv, otherwise skip it
            if Path(file).suffix.lower() != ".csv":
                continue

            full_path = os.path.join(filepath, file)
            df = pd.read_csv(full_path, encoding=encoding)
            df = df.loc[:, ~df.columns.str.contains("^Unnamed")]

            df_dict[Path(file).stem] = df

        question_state = cls(
            insights=df_dict["insights"],
            full_text=df_dict.get("full_text", None),
            chunks=df_dict.get("chunks", None)
        )
        
        # Normalize the columns - first get arrays to lists then get viable lists
        return question_state

    # ---------------------------------------------------------------------- #
    #                             HELPER UTILS                              #
    # ---------------------------------------------------------------------- #
    def _drop_unnamed_columns(self) -> None:
        for key, value in self.__dict__.items():
            if isinstance(value, pd.DataFrame):
                clean_value = value.loc[:, ~value.columns.str.contains("^Unnamed")]
                self.__dict__[key] = clean_value
    
    @staticmethod
    def _strict_literal_eval(value):
        if pd.isna(value):
            return []
        try:
            return ast.literal_eval(value)
        except (ValueError, TypeError, SyntaxError) as e:
            raise ValueError(
                f"Fatal Error: Failed to evaluate literal in 'paper_author' column. "
                f"Ensure ALL entries are strictly formatted as a Python list of strings, "
                f"e.g., ['Author A', 'Author B']. The offending value was: '{value}'. "
                f"Original error: {e.__class__.__name__}"
            ) from e

    def arrays_to_lists(self, columns):
        for attr_name, attr_value in self.__dict__.items():
            if isinstance(attr_value, pd.DataFrame):
                for col in columns:
                    if col in attr_value.columns:
                        attr_value[col] = attr_value[col].apply(
                            lambda v: v.tolist() if isinstance(v, np.ndarray) else v
                        )
                setattr(self, attr_name, attr_value)

    def normalize_list_columns(self, columns):
        for attr_name, attr_value in self.__dict__.items():
            if isinstance(attr_value, pd.DataFrame):
                for col in columns:
                    if col in attr_value.columns:
                        attr_value[col] = attr_value[col].apply(utils.ensure_list_of_strings)
                setattr(self, attr_name, attr_value)
        

class Summaries:
    def __init__(
        self, 
        llm_client: Any,
        ai_model: str,
        summaries: Optional[pd.DataFrame] = None, 
        summaries_folder: str = os.path.join(SUMMARY_SAVE_LOCATION, "parquet"),
        summary_string: Optional[str] = None, 
        ai_peer_review: Optional[Dict] = None,
        output_save_location: str = os.path.join(SUMMARY_SAVE_LOCATION, "results")
        ):
        """
        Wrapper for a DataFrame of summaries and tools for AI-assisted peer review.
        
        Args:
            summaries: DataFrame of either the summaries or the summaries synthesized into themes
            llm_client: LLM client for interacting with the AI model.
            ai_model: Name of the deep research model to use.
            summary_string: Optional; pre-computed concatenated summary string.
            ai_peer_review: Optional; stores the AI peer review output as a dictionary.
            output_save_location: Directory to save Word documents.
        """
        
        
        # If summaries is not provided, try and load from parquet
        if summaries is None:
            self.summaries = self.from_parquet(summaries_folder=summaries_folder, llm_client=llm_client, ai_model=ai_model)


        self.summaries: pd.DataFrame = summaries
        self.llm_client = llm_client
        self.ai_model = ai_model
        self.summary_string: Optional[str] = summary_string
        self.ai_peer_review: Optional[Dict] = ai_peer_review
        self.summaries_folder: str = summaries_folder
        self.output_save_location: str = output_save_location

        
    @classmethod
    def from_parquet(cls, summaries_folder: str, llm_client: Any, ai_model: str) -> "Summaries":
        """
        Load summaries from a parquet file containing a DataFrame.

        Args:
            filepath: Path to the parquet file.

        Returns:
            Summaries instance with loaded DataFrame.
        """

        files = os.listdir(cls.output_save_location)
        summaries_file = [file for file in files if file in ["summaries.parquet", "clean_summaries.parquet"]]
        if len(summaries_file) == 0:
            raise FileNotFoundError(f"No 'summaries.parquet' or 'clean_summaries.parquet' file found in {cls.output_save_location}, cannot load Summaries.")
        elif len(summaries_file) > 1:
            raise ValueError(f"Multiple summary files found in {cls.output_save_location}. Expected only one of 'summaries.parquet' or 'clean_summaries.parquet'. Either delete on file and retry, or load summary manually and pass to the constructor.")

        summaries = pd.read_parquet(os.path.join(summaries_folder, summaries_file[0]))

        return cls(summaries = summaries, 
                   llm_client = llm_client, 
                   ai_model = ai_model)

    def get_summary_string(self, output_result: bool = True) -> Optional[str]:
        """
        Concatenate cluster summaries into a single string per research question,
        ordered by question_id and cluster.

        Args:
            output_result: If True, return the concatenated string; else only sets self.summary_string.

        Returns:
            Concatenated summary string if output_result is True; else None.
        """
        # Ensure DataFrame is sorted by question_id and cluster for stable ordering
        self.summaries = self.summaries.sort_values(by=["question_id", "cluster"])

        output_string = ""
        for qid in self.summaries["question_id"].unique():
            qtext = self.summaries.loc[self.summaries["question_id"] == qid, "question_text"].iloc[0]
            question_df = self.summaries[self.summaries["question_id"] == qid]

            question_string = (
                f"Research question id: {qid}\n"
                f"Research question text: {qtext}\n"
                "Review:\n"
                f"{'\n\n'.join(question_df['cluster_summary'].tolist())}\n\n"
            )
            output_string += question_string

        self.summary_string = output_string

        if output_result:
            return output_string
   
    def gen_executive_summary(self, token_length: int = 600) -> Optional[str]:
        if not hasattr(self, "summaries"):
            raise ValueError("summaries attribute not found.")

        df = self.summaries.copy()
        df = df.reset_index(drop=False).rename(columns={"index": "_row"})  # preserve original order
        themed = {"label", "contents"}.issubset(df.columns)

        parts: list[str] = []
        for qtext, qdf in df.groupby("question_text", sort=False):
            parts.append(f"Question: {qtext}\n")
            qdf = qdf.sort_values("_row")
            if themed:
                for _, r in qdf.iterrows():
                    label = (r.get("label") or "").strip()
                    content = (r.get("contents") or "").strip()
                    if not content:
                        continue
                    parts.append(
                        f"Theme: {label}\n"
                        f"{content}\n"
                        "--- END THEME ---\n"
                    )
            else:
                for _, r in qdf.iterrows():
                    summ = (r.get("summary") or "").strip()
                    if not summ:
                        continue
                    parts.append(f"{summ}\n")
            parts.append("=== END QUESTION ===\n")

        if not parts:
            return None

        all_text = "\n".join(parts).strip()

        sys_prompt = Prompts().exec_summary(token_length=token_length)
        resp = call_chat_completion(
            sys_prompt=sys_prompt,
            user_prompt=all_text,
            llm_client=self.llm_client,
            ai_model=self.ai_model,
            return_json=True,
            fall_back={"executive_summary": "", "title": ""},
        )

        summary = (resp.get("executive_summary") or "").strip()
        title = (resp.get("title") or "").strip()
        self.exec_summary = summary or None
        self.title = title or None
        return self.title, self.exec_summary


    def gen_question_summaries(self) -> Optional[pd.DataFrame]:
        # Preconditions
        if not hasattr(self, "summaries"):
            return None
        df = self.summaries
        if "question_text" not in df.columns:
            raise ValueError("summaries must contain 'question_text'")

        # Decide input mode
        has_themes = {"label", "contents"}.issubset(df.columns)
        has_raw    = "summary" in df.columns
        if not (has_themes or has_raw):
            raise ValueError("summaries must have either ['label','contents'] or ['summary']")

        out_rows = []

        for qtext, qdf in df.groupby("question_text", sort=False):
            # Build payload for this question
            parts = [f"Question: {qtext}", ""]
            if has_themes:
                for _, r in qdf.iterrows():
                    parts.append(f"Theme: {r['label']}\n{r['contents']}\n")
            else:
                for _, r in qdf.iterrows():
                    parts.append(f"{r['summary']}\n")
            payload = "\n".join(parts).strip()

            # Call LLM
            sys_prompt = Prompts().question_summaries()
            fall_back = {"summary": ""}

            resp = call_chat_completion(
                sys_prompt=sys_prompt,
                user_prompt=payload,
                llm_client=self.llm_client,
                ai_model=self.ai_model,
                return_json=True,
                fall_back=fall_back,
            )

            out_rows.append({
                "question_text": qtext,
                "question_summary": (resp.get("summary") or "").strip(),
            })

        if not out_rows:
            return None

        qsum = (
            pd.DataFrame(out_rows)
            .drop_duplicates(subset=["question_text"], keep="last")
        )

        # One summary per question_text
        self.summaries = self.summaries.merge(qsum, on="question_text", how="left", validate="m:1")
        return self.summaries


    def summary_to_doc(self, paper_title: str = None, summary_filename: str = None) -> str:
        
        def _sanitize(text: str) -> str:
            control_chars = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')
            if text is None:
                return ""
            # ensure str
            s = str(text)
            # normalize unicode
            s = unicodedata.normalize("NFC", s)
            # drop XML-illegal control chars (keep \t, \n, \r)
            s = control_chars.sub('', s)
            return s
        
        if summary_filename is None:
            summary_filename = "literature_review.docx"

        if paper_title is None and hasattr(self, "title"):
            paper_title = self.title
        elif paper_title is None:
            raise ValueError("Paper_title must be provided. Either pass it in summary_to_doc or call get_executive_summary() which will provide a title.")

        doc = Document()
        doc.add_heading(_sanitize(paper_title), level=0)

        # Executive summary
        if getattr(self, "exec_summary", None):
            doc.add_heading("Executive summary", level=1)
            for para in str(self.exec_summary).splitlines():
                p = _sanitize(para)
                if p.strip():
                    doc.add_paragraph(p)
            doc.add_page_break()
        else: 
            raise ValueError("exec_summary attribute not found. Please run gen_executive_summary() before exporting to doc.")

        df = self.summaries.copy().reset_index(drop=False).rename(columns={"index": "_row"})
        themed = {"label", "contents"}.issubset(df.columns)

        if themed:
            for question_text, qdf in df.groupby("question_text", sort=False):
                doc.add_heading(_sanitize(question_text), level=1)
                if "question_summary" in qdf.columns:
                    qs = (qdf["question_summary"].dropna().astype(str).iloc[0]
                        if not qdf["question_summary"].dropna().empty else "")
                    if qs.strip():
                        doc.add_paragraph(_sanitize(qs))
                qdf = qdf.sort_values("_row")
                for _, r in qdf.iterrows():
                    label = _sanitize(r.get("label", ""))
                    contents = _sanitize(r.get("contents", ""))
                    if label:
                        doc.add_heading(f"Theme: {label}", level=2)
                    if contents:
                        for para in contents.splitlines():
                            p = _sanitize(para)
                            if p.strip():
                                doc.add_paragraph(p)
                doc.add_page_break()
        else:
            for question_text, qdf in df.groupby("question_text", sort=False):
                doc.add_heading(_sanitize(question_text), level=1)
                if "question_summary" in qdf.columns:
                    qs = (qdf["question_summary"].dropna().astype(str).iloc[0]
                        if not qdf["question_summary"].dropna().empty else "")
                    if qs.strip():
                        doc.add_paragraph(_sanitize(qs))
                qdf = qdf.sort_values("_row")
                for _, r in qdf.iterrows():
                    summ = _sanitize(r.get("summary", ""))
                    if summ:
                        for para in summ.splitlines():
                            p = _sanitize(para)
                            if p.strip():
                                doc.add_paragraph(p)
                doc.add_page_break()

        
        # Save
        save_dir = self.output_save_location
        os.makedirs(save_dir, exist_ok=True)
        
        out_path = os.path.join(save_dir, summary_filename)
        doc.save(out_path)
        
        # Mark that doc has been generated - as flag for running the AI peer review (neccesary as all the steps for running the dog generation are needed for peer reviewn).
        self.doc = True

        return (
            f'Word doc of the literature review generated and save here: {out_path}'
        )
    
    def get_ai_peer_review(self, 
                           save_directory: str = None,
                           save_filename: str = "ai_peer_review.parquet",
                           output_length: int = 5000, 
                           max_tokens: int = 10000) -> pd.DataFrame:
        """
        Request an AI peer review of the concatenated summaries.

        Args:
            output_length: Suggested maximum word count for the review.
            max_tokens: Hard limit on token usage for the AI model.

        Returns:
            AI review as a dataframe. If the review exceeds token budget, 'error' key may appear.
        """        
        # Function specific utilities ---------------------
        # Get the RQ and summaries excluding the current RQ
        def get_rq_summaries(self, current_rq_id) -> str:
            output_parts = []
            for rq_id, rq_df in self.summaries.groupby("question_id", sort=False):
                if rq_id == current_rq_id:
                    continue
                else:
                    output_parts.append(f"RESEARCH QUESTION:\n{rq_id}\n")
                    question_summary = rq_df["question_summary"].iloc[0]
                    output_parts.append(f"SUMMARY:\n{question_summary}\n")
            return "\n".join(output_parts)
        
        # Get the RQ and its associated content
        def get_rq_content(self, current_rq_id) -> str:
            current_rq_df = self.summaries[self.summaries["question_id"] == current_rq_id]
            current_rq_text = current_rq_df["question_text"].iloc[0]
            output_parts = []
            output_parts.append(f"RESEARCH QUESTION:\n{current_rq_text}\n")
            themed = {"label", "contents"}.issubset(self.summaries.columns)
            if themed:
                for _, row in current_rq_df.iterrows():
                    label = row["label"].strip()
                    content = row["contents"].strip()
                    output_parts.append(
                        f"Theme: {label}\n"
                        f"Summary: {content}\n"
                    )

                return "\n".join(output_parts)
            
            else:
                for _, row in current_rq_df.iterrows():
                    summary = row["summary"].strip()
                    output_parts.append(f"Summary: {summary}\n")
                return "\n".join(output_parts)
        
        # END UTILS -------------------------------

        # Check that executive summary has been generated
        if getattr(self, "exec_summary", None) is None:
            raise ValueError("Please run gen_executive_summary() before requesting an AI peer review.")
        
        # Check that the peer review has not already been generated
        if save_directory is None:
            save_directory = self.summaries_folder

        if os.path.exists(os.path.join(save_directory, save_filename)):
            recover = None
            while recover not in ["r", "n"]:
                recover = input("AI peer review file already exists. Loading existing file. (r)eload or create (n)ew? ")
            if recover == "r":
                self.ai_peer_review = pd.read_parquet(os.path.join(save_directory, save_filename))
                return self.ai_peer_review
            else:
                print("Generating new AI peer review and overwriting existing file.")

        # Populate the list that will form the prompt for the LLM
        output_list = []
        for rq_id, rq_df in self.summaries.groupby("question_id", sort=False):
            output_parts = []
            output_parts.append("INFORMATION FOR CONTEXT\n")
            output_parts.append("EXECUTIVE SUMMARY:\n")
            output_parts.append(self.exec_summary + "\n")
            output_parts.append("OTHER RESEARCH QUESTIONS AND SUMMARIES:\n")
            output_parts.append(get_rq_summaries(self, current_rq_id=rq_id))
            output_parts.append("CURRENT RESEARCH QUESTION AND CONTENT:\n")
            output_parts.append(get_rq_content(self, current_rq_id=rq_id))

            output_string = "\n".join(output_parts)
            
            # Call the reasoning model
            resp = call_reasoning_model(
                prompt=Prompts().ai_peer_review(
                    lit_review=output_string,
                    output_length=output_length,
                    max_tokens=max_tokens,
                    themed={"label", "contents"}.issubset(self.summaries.columns)
                ),
                llm_client=self.llm_client,
                ai_model=self.ai_model,
                timeout=1200
            )

            # Handle the response with a retry to clean the JSON via an LLM
            try:
                response_dict = json.loads(resp)
            except json.JSONDecodeError:
                response_retry = llm_json_clean(x = resp,
                                                prompt = Prompts().peer_review_format_check(),
                                                llm_client=self.llm_client,
                                                ai_model=self.ai_model)
                response_dict = json.loads(response_retry)

            overall_comment_df = pd.DataFrame({"comment_id": "overall", 
                                               "comment": response_dict.get("overall_comment"),
                                               "severity": pd.NA, 
                                               "location": "overall"}, index=[0])
                 

            specific_comments_df = pd.DataFrame(
                response_dict.get("specific_comments", []), 
                columns=["comment_id", "comment", "severity", "location"]
            )


            output = pd.concat([overall_comment_df, specific_comments_df], ignore_index=True)
            output["research_question_id"] = rq_id
            output["research_question_text"] = rq_df["question_text"].iloc[0]
            output["resubmit"] = response_dict.get("resubmit", False)

            output_list.append(output)

        output_df = pd.concat(output_list).reset_index(drop=True)
        # Convert to string and fill missing severity with "None" so that parquet will accept it
        output_df["severity"] = output_df["severity"].apply(lambda x: x if pd.notna(x) else "None")
        # Also make all id str (handle int and "overall" fields)
        output_df["comment_id"] = output_df["comment_id"].astype(str)

        self.ai_peer_review = output_df

        os.makedirs(save_directory, exist_ok=True)
        self.ai_peer_review.to_parquet(os.path.join(save_directory, save_filename), index=False)
        
        return self.ai_peer_review
    
    def peer_review_to_doc(self):
        if self.ai_peer_review is None:
            print("No AI peer review data available. Please run get_ai_peer_review() first.")
            return None
        
        # Convert the AI peer review DataFrame to a formatted string or document
        # This is a placeholder for actual document generation logic
        doc = Document()
        for _, row in self.ai_peer_review.iterrows():
            doc.add_paragraph(f"Comment ID: {row['comment_id']}")
            if row["severity"] is not None:
                doc.add_paragraph(f"Severity: {row['severity']}")
            if row["resubmit"]:
                doc.add_paragraph("**Resubmission Required**")
            if row["location"] == "overall":
                doc.add_paragraph(f"Location: {row['research_question_text']}")
            else:
                doc.add_paragraph(f"Location: {row['research_question_text']} - {row['location']}")
            doc.add_paragraph(f"Comment: {row['comment']}")
            doc.add_paragraph("---------------------")


        # Save
        save_dir = self.output_save_location
        os.makedirs(save_dir, exist_ok=True)

        out_path = os.path.join(save_dir, "peer_review.docx")
        doc.save(out_path)

        print(f"Peer review generated and saved to {out_path}")
        


