"""
Rendering and presentation layer for ReadingMachine.

This module converts completed synthesis artifacts into presentation-ready
outputs.

ReadingMachine distinguishes between:

    analytical state
        CorpusState + SummaryState

and

    render artifacts
        presentation-layer outputs

The analytical pipeline is responsible for reading, organizing, and
synthesizing corpus content. The rendering layer is responsible for
transforming those synthesis artifacts into formats suitable for human
consumption.

Typical rendering operations include:

    - stylistic rewriting
    - question-level summaries
    - executive-summary generation
    - document assembly
    - export to Markdown, DOCX, and PDF

Unlike the synthesis workflow, rendering does not create new analytical
structure. It operates on existing synthesis artifacts and produces
derived presentation outputs.

Render artifacts are persisted independently of SummaryState so that
expensive formatting operations can be reused without rerunning thematic
synthesis.

The module also provides traceability utilities that allow rendered
claims to be traced back through:

    theme
        ↓
    mapped insight
        ↓
    source chunk
        ↓
    source document

This preserves ReadingMachine's commitment to inspectability and
auditability even after narrative rendering has occurred.
"""

# Import custom libraries and modules
from . import config, utils
from .prompts import Prompts
from .state import SummaryState, CorpusState


# Import standard libraries
from typing import Any, Dict, Optional
import os
import pandas as pd
import unicodedata
import re
from docx import Document
import json
import hashlib
from pathlib import Path
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
import shutil


class Render:
    """
    Presentation and export layer for ReadingMachine synthesis outputs.

    The Render class converts completed synthesis artifacts into
    presentation-ready documents and provides optional rendering operations
    that improve readability without altering the underlying analytical
    outputs.

    Unlike the reading and synthesis stages, rendering does not generate
    new analytical structure. It operates on completed synthesis artifacts
    stored in SummaryState and produces derived presentation artifacts such
    as question overviews, executive summaries, stylistic rewrites, and
    exportable documents.

    Rendering therefore occupies a distinct layer of the ReadingMachine
    architecture:

        corpus reading
            ↓
        thematic synthesis
            ↓
        rendering
            ↓
        exported documents

    Render Sources
    --------------
    Rendering operates on a selected synthesis artifact drawn from
    SummaryState.

    Under normal operation, rendering begins from one of two terminal
    synthesis outputs:

    Orphan-corrected themes
        Fidelity-first output emphasizing coverage and completeness.

    Redundancy-reduced themes
        Readability-first output emphasizing reduced repetition across
        themes.

    Earlier synthesis artifacts can be rendered for diagnostic purposes
    when force rendering is enabled.

    Rendering Artifacts
    -------------------
    The Render class supports several optional presentation-layer
    transformations:

    Stylistic Rewrites
        Rewrites theme summaries into a specified narrative style while
        attempting to preserve analytical content.

    Question Summaries
        Generates concise narrative overviews for individual research
        questions.

    Executive Summary
        Produces a corpus-level synthesis and document title.

    Document Assembly
        Combines all rendering artifacts into a final render DataFrame that
        can be exported to external formats.

    Export
        Converts the assembled document into Markdown, DOCX, or PDF.

    Rendering Workflow
    ------------------
    A typical rendering workflow proceeds as follows:

        synthesis output
            ↓
        optional stylistic rewrite
            ↓
        question summaries
            ↓
        executive summary
            ↓
        document assembly
            ↓
        export

    Users may perform only a subset of these operations depending on their
    requirements.

    Render-State Model
    ------------------
    Render artifacts are intentionally separated from analytical state.

    ReadingMachine distinguishes between:

    Analytical State
        Stored in CorpusState and SummaryState.

    Render Artifacts
        Derived presentation outputs generated from analytical state.

    This separation allows expensive rendering operations to be reused
    without affecting the underlying synthesis outputs.

    Hash-Based Consistency
    ----------------------
    Render artifacts are protected by a deterministic summary-hash
    mechanism.

    When rendering begins, the selected synthesis artifact is normalized and
    hashed. Previously generated render artifacts are reused only when the
    stored hash matches the current synthesis content.

    This prevents presentation artifacts from being accidentally applied to
    a different synthesis state.

    Traceability
    ------------
    Although rendering focuses on presentation, traceability is preserved.

    The class includes utilities that allow users to trace rendered claims
    back through:

        theme
            ↓
        mapped insight
            ↓
        source chunk
            ↓
        source document

    This allows rendered outputs to remain connected to the underlying
    corpus despite multiple stages of synthesis and formatting.

    Design Principles
    -----------------
    The rendering layer follows several principles:

    Separation of Analysis and Presentation
        Rendering should not modify analytical outputs.

    Reusability
        Expensive rendering artifacts should be recoverable across
        sessions.

    Fidelity Preservation
        Rendering operations should preserve analytical content whenever
        possible.

    Traceability
        Rendered claims should remain linked to underlying evidence.

    Deterministic Recovery
        Previously generated render artifacts should only be reused when
        they correspond to the current synthesis state.

    Attributes
    ----------
    summary_state : SummaryState
        Synthesis-state object providing the source material for rendering.

    corpus_state : CorpusState
        Corpus-state object used for traceability operations.

    summary_to_render : pd.DataFrame
        Selected synthesis artifact used as the render source.

    render_path : str
        Directory used to store rendering artifacts.

    output_save_location : str
        Directory used for exported documents.

    render_object : tuple[str, int] or None
        Explicit render target selected by the user.

    force : bool
        Indicates whether rendering safeguards have been overridden.

    llm_client : Any
        Client used for rendering-stage LLM calls.

    ai_model : str
        Model identifier used for rendering operations.
    """
    def __init__(
        self,
        llm_client: Any,
        ai_model: str,
        summary_state: SummaryState = None,
        corpus_state: CorpusState = None,
        force: bool = False,
        render_object: tuple[str, int] = None,
        render_path: str = config.RENDER_SAVE_LOCATION,
        render_hash=config.summary_hash,
        output_save_location: str = os.path.join(config.OUTPUT_SAVE_LOCATION)
    ):
        """
        Initialize the rendering layer for synthesized ReadingMachine outputs.

        Prepares a Render object for converting synthesis artifacts stored in
        SummaryState into presentation-ready outputs. Rendering is distinct from
        the analytical pipeline: it operates on completed or explicitly selected
        synthesis artifacts and produces terminal presentation artifacts such as
        rewritten summaries, question overviews, executive summaries, titles, and
        exported documents.

        Initialization performs the following steps:

        1. Validate that SummaryState contains renderable synthesis artifacts.
        2. Validate any explicitly requested `render_object`.
        3. Select the summary artifact to render.
        4. Sort the selected summaries into canonical question/theme order.
        5. Compute a deterministic hash of the selected render input.
        6. Reload or reset previous render artifacts based on hash compatibility.

        Parameters
        ----------
        llm_client : Any
            Client used for rendering-related LLM calls, including stylistic
            rewriting, question summaries, executive summaries, and title
            generation.

        ai_model : str
            Model identifier used for rendering-stage LLM calls.

        summary_state : SummaryState, optional
            SummaryState containing synthesis artifacts generated by the
            Summarize pipeline.

        corpus_state : CorpusState, optional
            CorpusState used for traceability operations that link rendered text
            back to insights, chunks, and source documents.

        force : bool, default=False
            If True, allow rendering from explicitly selected or earlier-stage
            synthesis artifacts instead of enforcing the usual completed-stage
            selection logic.

        render_object : tuple[str, int], optional
            Explicit render target in the form `(artifact_name, index)`.

            Supported artifact names are:

            - `"cluster_summary"`
            - `"populated_theme"`
            - `"redundancy"`

            This option is intended for forced or diagnostic rendering.

        render_path : str, default=config.RENDER_SAVE_LOCATION
            Directory used to persist render-stage artifacts.

        render_hash : str, default=config.summary_hash
            Filename or key used to store the hash of the summary artifact being
            rendered.

        output_save_location : str, default=os.path.join(config.OUTPUT_SAVE_LOCATION)
            Directory where exported output files are written.

        Raises
        ------
        ValueError
            If SummaryState does not contain any cluster summaries or populated
            themes.

        ValueError
            If `render_object` is supplied but is not a two-element tuple.

        ValueError
            If `render_object` specifies an unsupported artifact name.

        Side Effects
        ------------
        May create or inspect files under `render_path`.

        May reload previously generated rendering artifacts when their stored
        hash matches the current summary input.

        May clear previous rendering artifacts through `_reinitialize_render()`
        if the stored hash does not match and the user elects to restart
        rendering.

        Attributes
        ----------
        summary_state : SummaryState
            Synthesis state used as the source for rendering.

        corpus_state : CorpusState
            Corpus state used for traceability and source lookup operations.

        summary_to_render : pd.DataFrame
            Selected synthesis artifact, sorted by `question_id` and
            `theme_id`, used as the render input.

        render_object : tuple[str, int] or None
            Explicit render target requested by the caller, if any.

        force : bool
            Whether force-rendering behavior is enabled.

        llm_client : Any
            LLM client used by rendering operations.

        ai_model : str
            Model identifier used by rendering operations.

        render_hash : str
            Hash filename or key used to validate render artifact reuse.

        render_path : str
            Directory used for render-stage persistence.

        output_save_location : str
            Directory used for final exported files.

        Notes
        -----
        Render artifacts are treated as presentation-layer outputs rather than
        analytical state. They are derived from SummaryState but are persisted
        separately so expensive formatting operations can be reused when the
        underlying synthesis artifact has not changed.

        The summary hash guards against accidental reuse of rendering artifacts
        generated from a different synthesis state.
        """
        # Check that we can instantiate this class - i.e. summaries have been generated
        if summary_state.cluster_summary_list == [] and summary_state.populated_theme_list == []:
            raise ValueError("SummaryState appears to be empty. Please run the summarization process before instantiating the Render class.")
        
        if render_object is not None:
            if not isinstance(render_object, tuple) or len(render_object) != 2:
                raise ValueError("render_object must be a tuple of length 2.")
            if render_object[0] not in ["cluster_summary", "populated_theme", "redundancy"]:
                raise ValueError(f"Invalid render_object '{render_object[0]}'. Valid options are 'cluster_summary', 'populated_theme', or 'redundancy'.")

        self.summary_state = summary_state
        self.corpus_state = corpus_state
        self.render_object = render_object
        self.force = force
        self.llm_client = llm_client
        self.ai_model = ai_model
        self.render_hash = render_hash
        self.render_path = render_path
        self.output_save_location: str = output_save_location

        # Get the exact content that we want to render on, depending on whether the use ran the redundancy pass or stopped at the final orphan pass ()
        self.summary_to_render = self._get_summaries_for_render(force=force)

        # Ensure sort for subsequent operations
        self.summary_to_render = self.summary_to_render.sort_values(by=["question_id", "theme_id"]).reset_index(drop=True)
        # Hash the summary to eiher check against previous renders or to set if this is the first time. This guards against changes in summary objects during rendering pass
        current_summary_hash = self._compute_df_hash(self.summary_to_render)

        # Handle check and recovery of pervious renders
        self._reinitialize_render(current_summary_hash=current_summary_hash)

    def _get_summaries_for_render(self, force) -> pd.DataFrame:
        """
        Select the synthesis artifact to use as the render source.

        Determines which SummaryState artifact should be rendered based on the
        current synthesis stage, the optional explicit `render_object`, and the
        `force` flag.

        Rendering normally uses one of two terminal synthesis artifacts:

        - redundancy-reduced themes, when the redundancy pass has completed
        - orphan-corrected populated themes, when orphan handling has completed
        but redundancy reduction has not been run

        Earlier-stage artifacts can be rendered only when `force=True`.

        Parameters
        ----------
        force : bool
            If True, allow rendering from explicitly selected or earlier-stage
            synthesis artifacts. If False, rendering is restricted to completed
            orphan or redundancy outputs.

        Returns
        -------
        pd.DataFrame
            Summary DataFrame selected for rendering.

        Raises
        ------
        ValueError
            If `render_object` is supplied while `force=False`.

        ValueError
            If the requested `render_object` index is out of range for the
            corresponding SummaryState artifact list.

        ValueError
            If the current summary state has not reached the orphan or redundancy
            stage and `force=False`.

        Notes
        -----
        If `self.render_object` is supplied, it must identify a renderable
        SummaryState artifact and index. Supported artifact names are:

        - `"cluster_summary"`
        - `"populated_theme"`
        - `"redundancy"`

        Automatic selection is based on `SummaryState.status(diagnostic=True)`.

        Redundancy outputs are treated as readability-first render sources because
        repeated content has been reduced, though citation or insight omission risk
        may be higher than immediately after orphan handling.

        Orphan-corrected populated themes are treated as fidelity-first render
        sources because mapped insights have been checked for coverage, but
        redundancy may remain.

        Force rendering is intended for diagnostics and intermediate inspection.
        It may render summaries that still contain unresolved orphans,
        redundancies, or incomplete synthesis artifacts.
        """
        # First check whether the user has specified a render object, if so make sure they have set flag to true
        if self.render_object is not None and force == False:
            raise ValueError("You have specified a render_object but force is set to False. Please set force to True to render the specified object, or set render_object to None to automatically determine which summaries to render based on the current stage of the summarization process.")
        # If render object and force flag are set return the specified object at the specified stage, make sure index is within range
        if self.render_object is not None and force == True:
            obj_name, index = self.render_object
            if obj_name == "cluster_summary":
                if not (-len(self.summary_state.cluster_summary_list) <= index < len(self.summary_state.cluster_summary_list)):
                    raise ValueError(f"Index {index} out of range for cluster_summary_list with length {len(self.summary_state.cluster_summary_list)}.")
                return self.summary_state.cluster_summary_list[index]
            if obj_name == "populated_theme":
                if not (-len(self.summary_state.populated_theme_list) <= index < len(self.summary_state.populated_theme_list)):
                    raise ValueError(f"Index {index} out of range for populated_theme_list with length {len(self.summary_state.populated_theme_list)}.")
                return self.summary_state.populated_theme_list[index]
            else:
                obj_name = "redundancy" # This is not needed, just here for explicit clarity
                if not (-len(self.summary_state.redundancy_list) <= index < len(self.summary_state.redundancy_list)):
                    raise ValueError(f"Index {index} out of range for redundancy_list with length {len(self.summary_state.redundancy_list)}.")
                return self.summary_state.redundancy_list[index]


        # If render object not set, automatically get the summary object. First get the standard approach - redundancy or orphan - if they are complete. Otherwise check the flaf force option and return other options
        status_diagnostic = self.summary_state.status(diagnostic=True)
        max_stage, last_stage = status_diagnostic["max_stage"], status_diagnostic["stage"]
        # If the last stage is mapped themes or theme schema, we have to check whether a populated theme was generated - i.e. is max stage > 1 (i.e. a full pass has happened). If so the last available render is on the last populated_theme
        if last_stage in ["theme_schema", "mapped_theme"] and max_stage > 1:
            last_stage = "populated_theme"
        elif last_stage in ["theme_schema", "mapped_theme", "cluster_summary"] and max_stage <= 1:
            last_stage = "cluster_summary"
        else:
            last_stage = last_stage
        
        
        if last_stage == "redundancy":
            print("Summaries have gone through the redundancy pass and been checked for orphans. This is a readability first output, but has a risk of insight ommission. If fidelity is your priority you should run the orphan pass again before rendering. SummaryState.handle_orphans().")
            return(self.summary_state.redundancy_list[0])
        if last_stage == "orphan":
            print("Summaries have gone through the orphan pass, but not been checked for redundancy. This is a fidelity first output and may contain redundancies. If you want to handle redundancies before rendering, please run the redundancy pass: SummaryState.address_redundancy().")
            return(self.summary_state.populated_theme_list[-1])
        if last_stage not in ["redundancy", "orphan"] and not force:
            raise ValueError(f"Your summary state is currently at stage '{last_stage}'. You can only render summaries if you have completed the orphan pass and/or the redundancy pass. Use SummaryState.status() to determine your current stage. " 
                                "If you know what you are doing and want to render with the current summaries, please instantiate the Render class with force=True to override this check.")
        if last_stage in ["cluster_summary", "populated_theme"] and force:
            print(f"WARNING Force rendering with summaries from stage '{last_stage}'. Be aware that the summaries may not be fully processed and may contain redundancies or unhandled orphans.")
            if last_stage == "populated_theme":
                return(self.summary_state.populated_theme_list[-1])
            else:
                return(self.summary_state.cluster_summary_list[0])
            
    @staticmethod
    def _compute_df_hash(df: pd.DataFrame) -> str:
        """
        Compute a deterministic content hash for a render source DataFrame.

        Produces a stable SHA-256 hash of a summary DataFrame after applying a
        series of normalization steps. The resulting hash is used to determine
        whether previously generated render artifacts remain valid for the
        current synthesis state.

        Parameters
        ----------
        df : pd.DataFrame
            DataFrame to hash.

        Returns
        -------
        str
            Deterministic SHA-256 hash of the normalized DataFrame.

            Returns `"EMPTY"` when the DataFrame is None or contains no rows.

        Notes
        -----
        Before hashing, the DataFrame is normalized by:

        - replacing missing values with a sentinel value
        - sorting columns alphabetically
        - sorting rows deterministically across all columns

        These normalization steps ensure that semantically identical DataFrames
        produce the same hash even if column order, row order, or missing-value
        representations differ.

        The normalized DataFrame is serialized to JSON before hashing.

        Within the rendering workflow, this hash is used to detect changes in
        the synthesis artifact being rendered. If the current hash differs from
        the stored render hash, previously generated render artifacts are
        assumed to be stale and may be discarded or regenerated.

        This method is intended for state-consistency checks rather than
        cryptographic security.
        """
        if df is None or df.empty:
            return "EMPTY"

        # Defensive copy
        df_normalized = df.copy()

        # Normalize NaNs
        df_normalized = df_normalized.fillna("__NULL__")

        # Sort columns alphabetically for consistency
        df_normalized = df_normalized.sort_index(axis=1)

        # Sort rows deterministically by all columns
        df_normalized = df_normalized.sort_values(
            by=list(df_normalized.columns)
        ).reset_index(drop=True)

        # Convert to stable JSON representation
        json_bytes = df_normalized.to_json(
            orient="records",
            date_format="iso"
        ).encode("utf-8")

        # Hash
        return hashlib.sha256(json_bytes).hexdigest()          
                   
    def _reinitialize_render(self, current_summary_hash: str) -> None:
        """
        Initialize, recover, or reset render-stage artifacts.

        Compares the hash of the current render source against any previously
        stored render hash and determines whether existing render artifacts can be
        safely reused.

        This method protects the rendering layer from accidentally reusing
        stylistic rewrites, question summaries, executive summaries, or assembled
        render outputs generated from a different synthesis artifact.

        Parameters
        ----------
        current_summary_hash : str
            Deterministic hash of `self.summary_to_render`, computed during
            initialization.

        Returns
        -------
        None

        Raises
        ------
        ValueError
            If a previous render hash exists, differs from the current summary
            hash, and the user chooses to abort rendering rather than overwrite
            previous render artifacts.

        Side Effects
        ------------
        Creates `self.render_path` if needed.

        May write a new render-hash Parquet file.

        May delete and recreate `self.render_path` if the user elects to restart
        rendering after a hash mismatch.

        May load existing render artifacts and attach them as attributes:

        - `title_exec_summary_df`
        - `question_summary_df`
        - `stylized_rewrite_df`

        May call `integrate_cosmetic_changes()` when all expected cosmetic
        artifacts are available.

        Notes
        -----
        The render hash is stored under:

            {self.render_path}/{self.render_hash}

        If no stored hash exists, the current hash is written and rendering starts
        from a clean artifact state.

        If the stored hash matches the current hash, compatible render artifacts
        are loaded from disk and reused.

        If the stored hash differs, the user must choose whether to abort or clear
        previous render artifacts. This prevents presentation-layer outputs from
        being silently applied to a different analytical summary state.

        Render artifacts are intentionally treated as derived presentation outputs,
        not analytical state. They may be regenerated without changing CorpusState
        or SummaryState.
        """
        # Check if any previous render objects exist:
        os.makedirs(self.render_path, exist_ok=True) # Ensure the render path exists

        if os.path.exists(os.path.join(self.render_path, self.render_hash)):
            saved_hash_df = pd.read_parquet(os.path.join(self.render_path, self.render_hash))
            saved_hash = saved_hash_df["summary_hash"].iloc[0]
            if saved_hash != current_summary_hash:
                restart = None
                while restart not in ["1", "2"]:
                    restart = input(
                        "A previous render hash was found that does not match the current summary hash. This suggests that the summaries have changed since the last render. " \
                        "Do you want to proceed with rendering and overwrite the previous render?\n"
                        "(1) Yes, proceed with rendering and overwrite the previous render.\n" 
                        "(2) No, stop the rendering process to avoid overwriting the previous render (this will abort render and require the correct summary_state be passed to Render init).\n"
                        "(1/2):\n"
                        ).strip()

                if restart == "2":
                    raise ValueError("Rendering process aborted by user to avoid overwriting previous render. Please check your summary_state and ensure it is correct before re-instantiating the Render class.")
                else:
                    # Clean up the render artifacts and start new render
                    shutil.rmtree(self.render_path) # This will remove all the previous render artifacts, which is important to avoid confusion and ensure a clean slate for the new render. We can be aggressive here with the cleanup because we have already confirmed with the user that they want to proceed with overwriting the previous render.
                    os.makedirs(self.render_path, exist_ok=True) # Re-create the render path after cleanup
                    # Create and save the new hash file
                    new_hash_df = pd.DataFrame({"summary_hash": [current_summary_hash]})
                    new_hash_df.to_parquet(os.path.join(self.render_path, self.render_hash), index=False)
                    print("Previous render artifacts deleted. Hash set to current summary_to_render.")
                    return(None)
                
            else: # if the hash matches we load the old dataframe and create all the attributes that we can from it.
                print("A previous render with the same summary hash was found. Loading the previously generated cosmetic artifacts. You can overwrite these by recalling the generation methods.\n")
                path = Path(self.render_path)
                loaded_artefacts = []
                for file in path.glob("*.parquet"):
                    if file.stem == config.render_prefix["render_title_exec_summary_df"]:
                        self.title_exec_summary_df = pd.read_parquet(file)
                        loaded_artefacts.append("title_summary_df")
                    elif file.stem == config.render_prefix["render_question_summary_df"]:
                        self.question_summary_df = pd.read_parquet(file)
                        loaded_artefacts.append("question_summary_df")
                    else:
                        if file.stem == config.render_prefix["render_stylized_rewrite_df"]:
                            self.stylized_rewrite_df = pd.read_parquet(file)
                            loaded_artefacts.append("stylized_rewrite_df")
                
                # If all three artefacts are loaded, we can integrate the cosmetic changes into a render_df that can be used for the final output generation. If some but not all cosmetic artefacts are loaded, we do not automatically integrate as the render_df would be incomplete, but we do inform the user of which artefacts were loaded and that they can be integrated by calling integrate_cosmetic_changes() once they have added any missing artefacts. If no cosmetic artefacts are loaded, we do not attempt to integrate and just inform the user that no cosmetic artefacts were found and they can proceed with generating new cosmetic artefacts or check their render save location if they believe they should be there.
                if len(loaded_artefacts) == 3:
                    final_render_df = self.integrate_cosmetic_changes() # I don't call render_df here (self.final_render_df gets set in integrate_cosmetic_changes), but mention it to be explicit
                    print("All cosmetic artefacts loaded and integrated. You can proceed with generating the final outputs.")
                
                if not hasattr(self, "final_render_df"):
                    print(f"{len(loaded_artefacts)} cosmetic artifacts loaded: {loaded_artefacts}. Final render_df not found. You should add any missing artifacts and/or call integrate_cosmetic_changes() to construct the final_render_df)")


        else: # if no previous file exists, save the hash and create the render df from the summary_to_render
            # Create and save the new hash file
            os.makedirs(self.render_path, exist_ok=True) # Ensure the render path exists
            new_hash_df = pd.DataFrame({"summary_hash": [current_summary_hash]})
            new_hash_df.to_parquet(os.path.join(self.render_path, self.render_hash), index=False)
            return(None)

    def stylistic_rewrite(self, style: str = "academic") -> pd.DataFrame:
        """
        Generate or reload stylistic rewrites of rendered theme summaries.

        Applies an optional presentation-layer rewrite to each thematic summary in
        `self.summary_to_render`. The rewrite changes tone, style, or register while
        attempting to preserve the analytical content of the underlying synthesis.

        Rewriting is performed sequentially within each research question. Previously
        rewritten themes for the same question are provided as frozen context so the
        model can maintain stylistic continuity across the rendered section.

        Parameters
        ----------
        style : str, default="academic"
            Target style passed to the rendering prompt. Interpretation of this
            value is controlled by `Prompts().stylistic_rewrite()`.

        Returns
        -------
        pd.DataFrame or None
            DataFrame containing:

            - `stylized_text`
            - `question_id`
            - `theme_id`

            Returns None if the user aborts because question-level summaries have
            already been generated.

        Side Effects
        ------------
        May prompt the user to reuse or regenerate existing stylistic rewrites.

        May prompt the user before proceeding if question-level summaries already
        exist.

        Mutates:

        - `self.stylized_rewrite_df`

        Persists the rewrite DataFrame to:

            {self.render_path}/{config.render_prefix["render_stylized_rewrite_df"]}.parquet

        Notes
        -----
        This is a rendering-layer operation, not an analytical synthesis step. It
        does not modify SummaryState, CorpusState, theme mappings, orphan outputs,
        or redundancy outputs.

        Stylistic rewriting is best performed before generating question-level
        summaries so the chosen register can propagate into later rendering
        artifacts.

        Because the operation rewrites synthesized text, it may introduce semantic
        drift or information loss. Users prioritizing maximum fidelity may choose to
        skip this step and render the original summaries directly.

        If an LLM call fails, the fallback records an empty rewritten summary for 
        that theme. The original analytical summary remains unchanged in `self.summary_to_render`.
        """
        # First check wheter a stylistic rewrite has already been done, if so offer reload or regenerate options
        if hasattr(self, "stylized_rewrite_df"): 
            rerun = None
            while rerun not in ["1", "2"]:
                rerun = input(
                    "A stylistic rewrite of the themes has already been generated. Do you want to regenerate the stylistic rewrite?\n"
                    "(1) Yes, regenerate the stylistic rewrite.\n" 
                    "(2) No, keep the existing stylistic rewrite.\n"
                    "(1/2):\n"
                ).strip()
                
            if rerun == "2":
                print("Keeping existing stylistic rewrite.")
                return self.stylized_rewrite_df
            else:                
                print("Regenerating stylistic rewrite.")
        
        # Then check whether the summaries have already been generated. If they have, send a warning that it is better to run the stylistic re-write prior to summarizing so that the tone is carried through
        if hasattr(self, "question_summary_df"):
            abort = None
            while abort not in ["1", "2"]:
                abort = input(
                    "You have already generated question summaries. It is generally recommended to do the stylistic rewrite before generating the question summaries so that the style is carried through to the question summaries. Do you want to proceed with the stylistic rewrite even though you have already generated question summaries?\n"
                    "(1) Yes, proceed with the stylistic rewrite even though I have already generated question summaries.\n" 
                    "(2) No, I want to do a stylistic rewrite before generating question summaries.\n"
                    "(1/2):\n"
                ).strip()

                if abort == "2":
                    print("Aborting stylistic rewrite so that you can do it before generating question summaries. Please run stylistic_rewrite() and then re-run gen_question_summaries() to generate the question summaries with the rewritten themes.")
                    return None
        
        stylized_content_df_out = pd.DataFrame(columns=["stylized_text", "question_id", "theme_id"])

        total_themes = self.summary_to_render.shape[0]
        count = 1

        for qid in self.summary_to_render["question_id"].unique():
            question_df = self.summary_to_render[self.summary_to_render["question_id"] == qid].copy()
            # Place frozen content here so that it resets for each rq
            frozen_content = ""
            for idx, row in question_df.iterrows():
                print(f"Stylistic rewrite for theme {count} of {total_themes}")
                question_text = row["question_text"]
                content = row["thematic_summary"]
                theme_label = row["theme_label"]
                theme_id = row["theme_id"]
                sys_prompt = Prompts().stylistic_rewrite(style=style, index=idx, label=theme_label)
                user_prompt = (
                    f"CURRENT RESEARCH QUESTION: {question_text}\n"
                    "FROZEN CONTENT:\n"
                    f"{frozen_content}\n\n"
                    f"CURRENT THEME LABEL: {theme_label}\n"
                    "THEMATIC SUMMARY TO STYLE:\n"
                    f"{content}\n\n"
                )
                fall_back = {"refined_summary": ""}

                json_schema = {
                    "name": "stylistic_rewrite_generator",
                    "strict": True,
                    "schema": {
                        "type": "object",
                        "properties": {
                            "refined_summary": {
                                "type": "string",
                                "description": "Refined thematic summary text."
                            }
                        },
                        "required": ["refined_summary"],
                        "additionalProperties": False
                    }
                }
                response = utils.call_chat_completion(
                    sys_prompt=sys_prompt,
                    user_prompt=user_prompt,
                    llm_client=self.llm_client,
                    ai_model=self.ai_model,
                    return_json=True,
                    json_schema=json_schema,
                    fall_back=fall_back
                )

                # Get the data, transform to df
                stylized_content = response["refined_summary"]
                stylized_content_df = pd.DataFrame({
                    "stylized_text": [stylized_content],
                    "question_id": [qid],
                    "theme_id": [theme_id],
                    })
                
                # Append to the output df
                stylized_content_df_out = pd.concat([stylized_content_df_out, stylized_content_df], ignore_index=True)
                frozen_content += f"{row['theme_label']}\n\n{stylized_content}\n\n"
                count += 1
        
        self.stylized_rewrite_df = stylized_content_df_out
        self.stylized_rewrite_df.to_parquet(os.path.join(self.render_path, config.render_prefix["render_stylized_rewrite_df"] + ".parquet"), index=False)
        return(stylized_content_df_out)
    
    def _gen_question_payload(self, qid) -> str:
        """
        Construct a structured theme-summary payload for a research question.

        Aggregates all theme summaries associated with a research question into a
        single formatted text block suitable for higher-level rendering tasks.

        The payload preserves the ordering of themes in `self.summary_to_render`
        and explicitly labels theme boundaries so that downstream LLM operations
        can distinguish individual thematic sections while reasoning over the
        question as a whole.

        Parameters
        ----------
        qid : str or int
            Research-question identifier.

        Returns
        -------
        str
            Concatenated thematic content in the form:

                Theme: <theme_label>
                <thematic_summary>
                --- END THEME ---

            repeated for each theme associated with the research question.

        Notes
        -----
        This payload is used as the primary input for higher-level rendering
        operations, including:

        - question-level summaries
        - executive summaries

        Theme ordering is preserved exactly as it appears in
        `self.summary_to_render`.

        The explicit theme boundary markers help reduce cross-theme blending and
        improve the model's ability to reason over the thematic structure of the
        synthesized output.
        """
        question_df = self.summary_to_render[self.summary_to_render["question_id"] == qid].copy()
        output = ""
        for _, row in question_df.iterrows():
            output += f"Theme: {row['theme_label']}\n"
            output += f"{row['thematic_summary']}\n"
            output += "--- END THEME ---\n"

        return(output)
    
    def gen_question_summaries(self) -> pd.DataFrame:
        """
        Generate or reload question-level overview summaries.

        Produces one concise narrative overview for each research question by
        synthesizing the theme summaries associated with that question. These
        overviews form an intermediate rendering layer between individual theme
        summaries and higher-level artifacts such as executive summaries or final
        document sections.

        Each research question is processed independently. Theme summaries are
        first assembled into a structured payload with explicit theme boundaries,
        then passed to the LLM for question-level synthesis.

        Returns
        -------
        pd.DataFrame or None
            DataFrame containing question-level summaries with columns:

            - `content`
            - `doc_attr`
            - `question_id`
            - `question_text`

            Returns None if the user aborts in order to run stylistic rewriting
            before question-summary generation.

        Side Effects
        ------------
        May prompt the user to reuse or regenerate existing question summaries.

        May prompt the user before proceeding if no stylistic rewrite exists.

        Mutates:

        - `self.question_summary_df`

        Persists question summaries to:

            {self.render_path}/{config.render_prefix["render_question_summary_df"]}.parquet

        Notes
        -----
        Question summaries are generated from `self.summary_to_render`, not directly
        from CorpusState or SummaryState.

        Stylistic rewriting, when used, is best performed before question-summary
        generation so the chosen tone can propagate into the question-level
        overviews.

        This method creates rendering artifacts only. It does not modify analytical
        synthesis state.

        If an LLM call fails, the fallback records an empty summary for that
        question while leaving the underlying theme summaries unchanged.
        """
        # First check whether question summaries already exist, if so offer reload or regenerate options
        if hasattr(self, "question_summary_df"):
            rerun = None
            while rerun not in ["1", "2"]:
                rerun = input(
                    "Question summaries already exist. Do you want to regenerate the question summaries?\n"
                    "(1) Yes, regenerate the question summaries.\n" 
                    "(2) No, keep the existing question summaries.\n"
                    "(1/2):\n"
                ).strip()
            if rerun == "2":
                print("Keeping existing question summaries.")
                return self.question_summary_df
            else:                
                print("Regenerating question summaries.")

        # Then check whether a stylistic re-write has been done. Flag for the user that if they want a stylistic re-write its likely better to do it after the summaries have been generated, but they can proceed as they like
        if not hasattr(self, "stylized_rewrite_df"):
            abort = None
            while abort not in ["1", "2"]:
                abort = input(
                    "You have not done a stylistic rewrite of the themes. This is not neccesary to proceed, "
                    "however if you intend to do a stylistic rewrite and want  the style of the re-write to appear in your summaries it is reccomended to do the re-write first. "
                    "Do you want to proceed with generating question summaries without doing a stylistic rewrite?\n"
                    "(1) Yes, proceed with generating question summaries without doing a stylistic rewrite.\n" 
                    "(2) No, I want to do a stylistic rewrite before generating question summaries.\n"
                    "(1/2):\n"
                ).strip()

                if abort == "2":
                    print("Aborting question summary generation so that you can do a stylistic rewrite first. Please run stylistic_rewrite() and then re-run gen_question_summaries() to generate the question summaries with the rewritten themes.")
                    return None
        
        # loop over question id to get the question payload
        question_summaries_df_list = []

        working_render_df = self.summary_to_render.copy()

        total_questions = working_render_df["question_id"].nunique()
        count = 1

        for qid in working_render_df["question_id"].unique():
            print(f"Generating summary for question {count} of {total_questions}")
            question_id = qid
            question_text = working_render_df[working_render_df["question_id"] == qid]["question_text"].iloc[0]
            question_payload = self._gen_question_payload(qid=question_id)

        # Prepare the full payload for the llm call
            sys_prompt = Prompts().question_summaries()
            user_prompt = (
            f"Research question: {question_text}\n\n"
            "CONTENT TO SUMMARIZE:\n"
            f"{question_payload}\n\n"
            )

            fall_back = {"summary": ""}
            json_schema = {
                "name": "question_summary_generator",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "summary": {
                            "type": "string",
                            "description": "A single-paragraph thematic overview (3-5 sentences) synthesizing the themes for the research question."
                        }
                    },
                    "required": ["summary"],
                    "additionalProperties": False
                }
            }
            
            # Call the llm
            response = utils.call_chat_completion(
                sys_prompt=sys_prompt,
                user_prompt=user_prompt,
                llm_client=self.llm_client,
                ai_model=self.ai_model,
                return_json=True,
                json_schema=json_schema,
                fall_back=fall_back
            )

            question_summary_txt = response["summary"]
            question_summary_df = pd.DataFrame({
                "content": [question_summary_txt],
                "doc_attr": ["question_summary"],
                "question_id": [qid],
            })

            question_summaries_df_list.append(question_summary_df)
            count += 1

        question_summaries_df_out = pd.concat(question_summaries_df_list, ignore_index=True).sort_values(by=["question_id"]).reset_index(drop=True)
        question_summaries_df_out = (
            question_summaries_df_out
            .merge(self.corpus_state.questions[["question_id", "question_text"]],
                   on="question_id",
                   how="left")
        )

        self.question_summary_df = question_summaries_df_out
        self.question_summary_df.to_parquet(os.path.join(self.render_path, config.render_prefix["render_question_summary_df"] + ".parquet"), index=False)       
        return(self.question_summary_df)
    
    
    def _gen_exec_summary_payload(self) -> str:
        """
        Construct the corpus-level payload used for executive-summary generation.

        Aggregates all rendered theme summaries into a single structured text
        representation covering every research question in the synthesis output.
        The payload preserves both question-level and theme-level boundaries so
        that the executive-summary generator can reason over the full thematic
        structure of the corpus.

        Returns
        -------
        str
            Structured text payload in the form:

                Research Question: <question_text>

                Theme: <theme_label>
                <thematic_summary>
                --- END THEME ---

                ...

                === END QUESTION ===

            repeated for each research question in the render source.

        Notes
        -----
        Theme ordering is preserved from `self.summary_to_render`.

        Question sections are constructed using `_gen_question_payload()` and then
        combined into a single corpus-level representation.

        This payload serves as the primary input for executive-summary generation,
        allowing the model to synthesize across research questions while retaining
        visibility into the underlying thematic structure.

        Explicit theme and question boundary markers help reduce unintended
        blending of concepts across themes and research questions during
        summarization.
        """
        temp_summary_df = self.summary_to_render.copy()
        temp_summary_df = temp_summary_df.sort_values(by=["question_id", "theme_id"]).reset_index(drop=True)

        # Loop over the questions
        output = ""
        for qid in temp_summary_df["question_id"].unique():
            output += f"Research Question: {temp_summary_df[temp_summary_df['question_id'] == qid]['question_text'].iloc[0]}\n"
            question_output = self._gen_question_payload(qid)
            output += question_output
            output += "\n=== END QUESTION ===\n\n"

        return output
    

    def gen_exec_summary(self, word_count: int = 500) -> pd.DataFrame:
        """
        Generate or reload the document title and executive summary.

        Produces corpus-level rendering artifacts from the full set of rendered
        theme summaries. The method constructs a structured payload covering all
        research questions and themes, then asks the LLM to generate both a concise
        document title and an executive summary.

        Parameters
        ----------
        word_count : int, default=500
            Target length, in words, for the executive summary.

        Returns
        -------
        pd.DataFrame
            DataFrame containing two rows:

            - document title
            - executive summary

            Columns:

            - `content`
            - `doc_attr`

        Side Effects
        ------------
        May prompt the user to reuse or regenerate an existing title and executive
        summary.

        Mutates:

        - `self.title_exec_summary_df`

        Persists the generated artifact to:

            {self.render_path}/{config.render_prefix["render_title_exec_summary_df"]}.parquet

        Notes
        -----
        The executive summary is generated from `self.summary_to_render` through
        `_gen_exec_summary_payload()`. It is a rendering-layer artifact and does
        not modify CorpusState, SummaryState, mappings, populated themes, or
        redundancy outputs.

        If an LLM call fails, the fallback records empty strings for both the title
        and executive summary while leaving the underlying rendered summaries
        unchanged.
        """

        if hasattr(self, "title_exec_summary_df"):
            rerun = None
            while rerun not in ["1", "2"]:
                rerun = input(
                    "An executive summary and title have already been generated. Do you want to regenerate the executive summary and title?\n"
                    "(1) Yes, regenerate the executive summary and title.\n" 
                    "(2) No, keep the existing executive summary and title.\n"
                    "(1/2):\n"
                ).strip()
                
            if rerun == "2":
                print("Keeping existing executive summary and title.")
                return self.title_exec_summary_df
            else:                
                print("Regenerating executive summary and title.")

        exec_summary_payload = self._gen_exec_summary_payload()
        
        sys_prompt = Prompts().exec_summary(word_count=word_count)
        user_prompt = exec_summary_payload
        fall_back = {"executive_summary": "", "title": ""}
        json_schema = {
            "name": "executive_summary_generator",
            "strict": True,
            "schema": {
                "type": "object",
                "properties": {
                    "executive_summary": {
                        "type": "string",
                        "description": "The final executive summary text in continuous prose."
                    },
                    "title": {
                        "type": "string",
                        "description": "A concise descriptive title, maximum 12 words, no subtitle."
                    }
                },
                "required": ["executive_summary", "title"],
                "additionalProperties": False
            }
        }

        response = utils.call_chat_completion(
            sys_prompt=sys_prompt,
            user_prompt=user_prompt,
            llm_client=self.llm_client,
            ai_model=self.ai_model,
            return_json=True,
            json_schema=json_schema,
            fall_back=fall_back
        )

        exec_summary = response["executive_summary"]
        title = response["title"]
        
        title_exec_summary_df = pd.DataFrame({
            "content": [title, exec_summary],
            "doc_attr": ["title", "exec_summary"]
            })

        self.title_exec_summary_df = title_exec_summary_df

        self.title_exec_summary_df.to_parquet(os.path.join(self.render_path, config.render_prefix["render_title_exec_summary_df"] + ".parquet"), index=False)

        return self.title_exec_summary_df

    
    def _add_question_summaries_to_render_df(
        self,
        render_df: pd.DataFrame,
        question_summaries: list
    ) -> pd.DataFrame:
        """
        Insert question-level summaries into a render DataFrame.

        Adds one question-summary row before the theme-summary rows associated
        with each research question. This creates the narrative hierarchy used by
        final rendering:

            question summary
                ↓
            theme summaries

        Parameters
        ----------
        render_df : pd.DataFrame
            Theme-level render DataFrame.

        question_summaries : list[str]
            Question-level summaries ordered to match
            `self.corpus_state.questions["question_id"]`.

        Returns
        -------
        pd.DataFrame
            Render DataFrame with question-summary rows inserted before their
            corresponding theme-summary rows.

        Notes
        -----
        The method assumes that `question_summaries` is ordered consistently with
        `self.corpus_state.questions`.

        Inserted question-summary rows use:

        - `content`
        - `doc_attr = "question_summary"`
        - `question_id`

        The method restores `question_text` by merging against
        `self.corpus_state.questions`.

        This helper only modifies the presentation-layer render structure. It does
        not generate summaries, modify analytical state, or add title/executive
        summary rows.
        """
        working_render_df = render_df.copy()

        # Separate the working_render_df into a list of dataframes, one for each question
        list_of_qdfs = []
        for qid in self.corpus_state.questions["question_id"].to_list():
            question_df = working_render_df[working_render_df["question_id"] == qid].copy()
            list_of_qdfs.append(question_df)
        
        # Concat the question summaries with the corresponding question dfs and create a new list of dataframes with summaries
        list_of_qdfs_with_summaries = []
        for qdf, q_summary in zip(list_of_qdfs, question_summaries):
            question_summary_df = pd.DataFrame({
                "content": [q_summary],
                "doc_attr": ["question_summary"],
                "question_id": [qdf["question_id"].iloc[0]]
            })
            qdf_with_summary = pd.concat([question_summary_df, qdf], ignore_index=True)
            list_of_qdfs_with_summaries.append(qdf_with_summary)

        # Concat the list of question dfs with summaries into a single dataframe
        questions_with_summaries_df = pd.concat(list_of_qdfs_with_summaries, ignore_index=True)
        questions_with_summaries_df = (
            questions_with_summaries_df
            .drop(columns=["question_text"], errors="ignore")
            .merge(self.corpus_state.questions[["question_id", "question_text"]],
                   on="question_id",
                   how="left")
        )
        # Return the resulting dataframe
        return questions_with_summaries_df


    def integrate_cosmetic_changes(self) -> pd.DataFrame:
        """
        Assemble the final render DataFrame from available render artifacts.

        Combines the selected synthesis summaries with optional presentation-layer
        artifacts generated during rendering, including question summaries,
        stylistic rewrites, title, and executive summary. The resulting DataFrame
        is the canonical input for final document export.

        Returns
        -------
        pd.DataFrame
            Final render DataFrame with document components ordered for export.

        Raises
        ------
        ValueError
            If `self.summary_to_render` no longer matches the render hash recorded
            during Render initialization.

        Side Effects
        ------------
        Mutates:

        - `self.final_render_df`

        Persists the final render DataFrame to:

            {self.render_path}/{config.render_prefix["final_render_df"]}.parquet

        Notes
        -----
        Hash validation ensures that presentation-layer artifacts are only applied
        to the same synthesis artifact from which they were generated.

        The assembly process proceeds as follows:

        1. Validate the current summary hash against the stored render hash.
        2. Convert theme summaries into base render rows.
        3. Insert question-summary rows when available.
        4. Merge stylistic rewrites when available.
        5. Prepend title and executive-summary rows when available.
        6. Assign deterministic `doc_order`.
        7. Drop synthesis diagnostics not needed for rendering.

        If stylistic rewrites exist, missing `stylized_text` values are filled
        from `content` so non-theme rows and unrewritten rows still have renderable
        text.

        This method assembles rendering artifacts only. It does not generate new
        LLM content and does not modify CorpusState or SummaryState.
        """

        # --- Hash validation ---
        current_summary_hash = self._compute_df_hash(self.summary_to_render)
        saved_hash_df = pd.read_parquet(os.path.join(self.render_path, self.render_hash))
        saved_hash = saved_hash_df["summary_hash"].iloc[0]

        if current_summary_hash != saved_hash:
            raise ValueError(
                "The summary_to_render has changed since Render initialization. "
                "Reinitialize Render before integrating artifacts."
            )

        # --- Base render df ---
        render_df = (
            self.summary_to_render
            .sort_values(by=["question_id", "theme_id"])
            .copy()
        )

        render_df.rename(columns={"thematic_summary": "content"}, inplace=True)
        render_df["doc_attr"] = "thematic_summary"

        # --- Question summaries ---
        if hasattr(self, "question_summary_df"):
            render_df = self._add_question_summaries_to_render_df(
                render_df,
                self.question_summary_df["content"].to_list()
            )

        # --- Stylized rewrite ---
        if hasattr(self, "stylized_rewrite_df"):
            render_df = render_df.merge(
                self.stylized_rewrite_df[["question_id", "theme_id", "stylized_text"]],
                on=["question_id", "theme_id"],
                how="left"
            )
            render_df["stylized_text"] = render_df["stylized_text"].fillna(render_df["content"])

        # --- Title + Exec summary ---
        if hasattr(self, "title_exec_summary_df"):
            title_exec_df = self.title_exec_summary_df.copy()
            # Check if there is a stylized rewrite, in which case copy the content column to the stylized text so we have something to render
            if "stylized_text" in render_df.columns:
                title_exec_df["stylized_text"] = title_exec_df["content"] # This is to ensure that if there is a stylized rewrite for the title and exec summary, which otherwise are under content.
            # use safe concat with schema to ensure that schema is adopted - to keep ordering dtypes correct, take schema from render.
            render_df = utils.concat_with_schema(title_exec_df, render_df, schema_from="bottom")



        # --- Final ordering ---
        render_df["doc_order"] = range(len(render_df))
        render_df.drop(columns = ["allocated_length", "current_length", "perc_of_max_length", "length_flag"], inplace=True, errors="ignore") # We don't need this column in the render df and it can cause issues if it is left in with null values, so we drop it if it exists. We set errors to ignore just in case it doesn't exist, which can happen if the summary_to_render was generated with an older version of the code that didn't have allocated_length in the summary df.

        # --- Save ---
        self.final_render_df = render_df
        self.final_render_df.to_parquet(
            os.path.join(self.render_path, config.render_prefix["final_render_df"] + ".parquet"),
            index=False
        )
        return self.final_render_df
    
    def render_output(
        self,
        output_type: str = "md",
        use_stylized: bool = False,
        filename: str = "rendered_output"
    ) -> None:
        """
        Export the assembled render DataFrame to a document.

        Converts `self.final_render_df` into a final presentation artifact and
        writes it to disk in the requested format. This is the terminal step of
        the rendering workflow.

        Parameters
        ----------
        output_type : str, default="md"
            Output format.

            Must be one of:

            - `"md"`
            - `"docx"`
            - `"pdf"`

        use_stylized : bool, default=False
            If True, render using `stylized_text` where available. Missing
            stylized values fall back to `content`.

            If False, render using the original rendered text stored in
            `content`.

        filename : str, default="rendered_output"
            Output filename without extension.

        Returns
        -------
        None

        Raises
        ------
        ValueError
            If `self.final_render_df` does not exist.

        ValueError
            If `output_type` is not one of `"md"`, `"docx"`, or `"pdf"`.

        Side Effects
        ------------
        Creates `self.output_save_location` if it does not already exist.

        Writes an output file to:

            {self.output_save_location}/{filename}.{output_type}

        Notes
        -----
        The method requires `integrate_cosmetic_changes()` to have been run so
        that `self.final_render_df` exists.

        A temporary `final_text` column is created to determine which text source
        should be rendered:

        - `stylized_text` when `use_stylized=True`
        - `content` otherwise

        Format-specific rendering is delegated to:

        - `_render_to_markdown()`
        - `_render_to_docx()`
        - `_render_to_pdf()`

        This method performs only output generation. It does not modify
        SummaryState, CorpusState, or any persisted rendering artifacts.
        """

        if not hasattr(self, "final_render_df"):
            raise ValueError(
                "final_render_df not found. Please run integrate_cosmetic_changes() first."
            )

        if output_type not in ["md", "docx", "pdf"]:
            raise ValueError("output_type must be one of: 'md', 'docx', 'pdf'.")
        
        os.makedirs(self.output_save_location, exist_ok=True)

        df = self.final_render_df.copy()

        # Choose text source
        if use_stylized and "stylized_text" in df.columns:
            df["final_text"] = df["stylized_text"].fillna(df["content"])
        else:
            df["final_text"] = df["content"]

        # Dispatch
        if output_type == "md":
            self._render_to_markdown(df, filename)

        elif output_type == "docx":
            self._render_to_docx(df, filename)

        elif output_type == "pdf":
            self._render_to_pdf(df, filename)

        print(f"Render complete: {filename}.{output_type}")


    def _render_to_markdown(self, df: pd.DataFrame, filename: str) -> None:
        """
        Render the final document as a Markdown file.

        Converts an ordered render DataFrame into a Markdown document using
        heading levels to represent the document hierarchy.

        Parameters
        ----------
        df : pd.DataFrame
            Render DataFrame containing ordered rows and a `final_text` column.

        filename : str
            Output filename without extension.

        Returns
        -------
        None

        Side Effects
        ------------
        Writes a Markdown file to:

            {self.output_save_location}/{filename}.md

        Notes
        -----
        Rows are rendered in `doc_order`.

        The document hierarchy is determined by `doc_attr`:

        - `title`: level-one heading
        - `exec_summary`: executive-summary section
        - `question_summary`: research-question heading and overview
        - `thematic_summary`: theme-level section

        This method performs format conversion only. It does not modify render
        artifacts, CorpusState, or SummaryState.
        """

        lines = []

        for _, row in df.sort_values("doc_order").iterrows():

            text = row["final_text"]
            attr = row["doc_attr"]
            question_text = row["question_text"]

            if attr == "title":
                lines.append(f"# {text}\n")

            elif attr == "exec_summary":
                lines.append(f"## EXECUTIVE SUMMARY\n{text}\n")

            elif attr == "question_summary":
                lines.append(f"## {question_text}\n## Question Overview\n{text}\n")

            elif attr == "thematic_summary":
                lines.append(f"### {row.get('theme_label','')}\n{text}\n")

        output_path = os.path.join(self.output_save_location, f"{filename}.md")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _render_to_docx(self, df: pd.DataFrame, filename: str) -> None:
        """
        Render the final document as a Microsoft Word file.

        Converts an ordered render DataFrame into a DOCX document using Word
        heading levels to represent the document hierarchy.

        Parameters
        ----------
        df : pd.DataFrame
            Render DataFrame containing ordered rows and a `final_text` column.

        filename : str
            Output filename without extension.

        Returns
        -------
        None

        Side Effects
        ------------
        Writes a DOCX file to:

            {self.output_save_location}/{filename}.docx

        Notes
        -----
        Rows are rendered in `doc_order`.

        The document hierarchy is determined by `doc_attr`:

        - `title`: document title
        - `exec_summary`: executive-summary section
        - `question_summary`: research-question heading and overview
        - `thematic_summary`: theme-level section

        This method performs format conversion only. It does not modify render
        artifacts, CorpusState, or SummaryState.
        """

        document = Document()

        for _, row in df.sort_values("doc_order").iterrows():

            text = row["final_text"]
            attr = row["doc_attr"]
            question_text = row["question_text"]

            if attr == "title":
                document.add_heading(text, level=0)

            elif attr == "exec_summary":
                document.add_heading("EXECUTIVE SUMMARY", level=1)
                document.add_paragraph(text)

            elif attr == "question_summary":
                document.add_heading(question_text, level=1)
                document.add_heading("Question Overview", level=1)
                document.add_paragraph(text)

            elif attr == "thematic_summary":
                document.add_heading(row.get("theme_label",""), level=2)
                document.add_paragraph(text)

        output_path = os.path.join(self.output_save_location, f"{filename}.docx")
        document.save(output_path)

    def _render_to_pdf(self, df: pd.DataFrame, filename: str) -> None:
        """
        Render the final document as a PDF file.

        Converts an ordered render DataFrame into a PDF document using ReportLab.
        Document structure is derived from row-level `doc_attr` values and rendered
        with heading and paragraph styles.

        Parameters
        ----------
        df : pd.DataFrame
            Render DataFrame containing ordered rows and a `final_text` column.

        filename : str
            Output filename without extension.

        Returns
        -------
        None

        Side Effects
        ------------
        Writes a PDF file to:

            {self.output_save_location}/{filename}.pdf

        Notes
        -----
        Rows are rendered in `doc_order`.

        Line breaks in generated text are normalized into ReportLab-compatible
        HTML-style break tags before rendering.

        The document hierarchy is determined by `doc_attr`:

        - `title`: document title
        - `exec_summary`: executive-summary section
        - `question_summary`: research-question heading and overview
        - `thematic_summary`: theme-level section

        This method performs format conversion only. It does not modify render
        artifacts, CorpusState, or SummaryState.
        """
        def normalize_for_pdf(text):
            # Convert double newlines into paragraph breaks
            text = re.sub(r"\n\s*\n", "<br/><br/>", text)
            # Convert remaining single newlines
            text = text.replace("\n", "<br/>")
            return text

        output_path = os.path.join(self.output_save_location, f"{filename}.pdf")
        doc = SimpleDocTemplate(output_path)

        styles = getSampleStyleSheet()
        elements = []

        for _, row in df.sort_values("doc_order").iterrows():

            text = row["final_text"]
            text = normalize_for_pdf(text)
            attr = row["doc_attr"]
            question_text = row["question_text"]

            if attr == "title":
                elements.append(Paragraph(f"<b>{text}</b>", styles["Title"]))
                elements.append(Spacer(1, 0.3 * inch))

            elif attr == "exec_summary":
                elements.append(Paragraph("<b>EXECUTIVE SUMMARY</b>", styles["Heading1"]))
                elements.append(Paragraph(text, styles["BodyText"]))
                elements.append(Spacer(1, 0.2 * inch))

            elif attr == "question_summary":
                elements.append(Paragraph(f"<b>{question_text}</b>", styles["Heading1"]))
                elements.append(Paragraph("<b>Question Overview</b>", styles["Heading1"]))
                elements.append(Paragraph(text, styles["BodyText"]))
                elements.append(Spacer(1, 0.2 * inch))

            elif attr == "thematic_summary":
                elements.append(Paragraph(f"<b>{row.get('theme_label','')}</b>", styles["Heading2"]))
                elements.append(Paragraph(text, styles["BodyText"]))
                elements.append(Spacer(1, 0.2 * inch))

        doc.build(elements)

    def trace_claim(
        self,
        question_text: str,
        theme_label: str,
        citation_lastname: list,
        citation_year: int
    ) -> pd.DataFrame:
        """
        Trace a rendered claim back to candidate source insights and chunks.

        Retrieves source evidence associated with a claim by following the
        ReadingMachine traceability chain from rendered theme context to mapped
        insights and then back to source document chunks.

        The lookup proceeds as follows:

            question text + theme label
                ↓
            question_id + theme_id
                ↓
            mapped insight IDs
                ↓
            corpus insight records
                ↓
            source chunks and document metadata

        Parameters
        ----------
        question_text : str
            Research-question text identifying the rendered section containing
            the claim.

        theme_label : str
            Theme label identifying the rendered theme containing the claim.

        citation_lastname : list[str]
            Author last names used to filter candidate source records. Names are
            combined into a regex pattern and matched against `paper_author`.

        citation_year : int
            Publication year used to filter candidate source records.

        Returns
        -------
        pd.DataFrame
            Candidate source records containing:

            - `insight_id`
            - `insight`
            - `chunk_id`
            - `chunk_text`
            - `paper_title`
            - `paper_author`
            - `paper_date`

        Notes
        -----
        This method provides a lightweight provenance lookup for rendered claims.
        It does not prove that a particular sentence was generated from a specific
        insight; instead, it narrows the search to insights mapped to the relevant
        theme and matching the supplied citation metadata.

        Matching is based on rendered question text, theme label, author-name
        regex matching, and publication year. If citation formatting changes or
        author names are ambiguous, returned candidates may need manual review.

        `self.final_render_df` must exist before this method is used.

        Future versions may support direct lookup by `insight_id` or structured
        citation identifiers, which would provide more precise traceability than
        citation metadata matching.
        """
        
        # Get the theme_id for the question and theme label
        question_id = self.final_render_df[self.final_render_df["question_text"] == question_text]["question_id"].iloc[0]
        theme_id = self.final_render_df[
            (self.final_render_df["question_id"] == question_id) & (self.final_render_df["theme_label"] == theme_label)
        ]["theme_id"].iloc[0]

        # use theme id to get the insights that were mapped to the theme in the summary state
        possible_insights_df = self.summary_state.mapped_theme_list[-1][
            (self.summary_state.mapped_theme_list[-1]["theme_id"] == theme_id)
            & (self.summary_state.mapped_theme_list[-1]["question_id"] == question_id)
        ].copy()

        possible_insights_lst = possible_insights_df["insight_id"].tolist()

        # Get all the insights that match the citation info and the theme_id, merge with chunks to get the chunk_text
        possible_insights_df = self.corpus_state.insights[self.corpus_state.insights["insight_id"].isin(possible_insights_lst)].copy()
        regex_pattern = "|".join(citation_lastname)
        possible_insights_df = (
            possible_insights_df[
            possible_insights_df["paper_author"].str.contains(
                pat = regex_pattern, case=False, na=False, regex=True
                ) 
                & (possible_insights_df["paper_date"] == citation_year)
            ].merge(
                self.corpus_state.chunks[["paper_id", "chunk_id", "chunk_text"]], 
                how="left",
                on=["chunk_id","paper_id"]
            ).copy()
        )
        # Return the relevant columns for the user to parse. Allows them to confirm they have the correct citation.
        possible_insights_df = possible_insights_df[["insight_id", "insight", "chunk_id", "chunk_text", "paper_title", "paper_author", "paper_date"]]
        return(possible_insights_df)
        