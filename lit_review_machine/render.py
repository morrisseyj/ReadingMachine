
# Import custom libraries and modules
from lit_review_machine import config, utils
from lit_review_machine.prompts import Prompts
from lit_review_machine.state import SummaryState, CorpusState


# Import standard libraries
from typing import Any, Dict, Optional
import os
import pandas as pd
import unicodedata
import re
from docx import Document
import json
import hashlib
from pathlib import Path
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch



class Render:
    def __init__(
        self, 
        llm_client: Any,
        ai_model: str,
        summary_state: SummaryState = None, 
        corpus_state: CorpusState = None,
        force: bool = False,
        render_object: tuple[str, int] = None,
        render_path: str = config.RENDER_SAVE_LOCATION,
        render_hash = config.summary_hash,
        output_save_location: str = os.path.join(config.SUMMARY_SAVE_LOCATION, "results")
        ):
        """
        Wrapper for a DataFrame of summaries and tools for AI-assisted peer review.
        
        Args:
            summary_state: SummaryState object containing the summaries and related metadata.
            corpus_state: CorpusState object containing the paper metadata and related information.
            force: If True allows rendering of summaries even if the summarization process has not been completed - either orphan handling or redundancy pass. Use with caution as the summaries may not be fully processed and may contain redundancies or unhandled orphans. 
            render_object: Optional tuple specifying which object to render and at which stage (articulated as list index) (e.g. "cluster_summary", 0" or "populated_theme", -1"). Must be passed with force=True. If None, the class will automatically determine which summaries to render based on the current stage of the summarization process as indicated by summary_state.status().
            summary_string: Optional; pre-computed concatenated summary string.
            ai_peer_review: Optional; stores the AI peer review output as a dictionary.
            output_save_location: Directory to save Word documents.
        """
        # Check that we can instantiate this class - i.e. summaries have been generated
        if summary_state.cluster_summary_list == [] and summary_state.populated_theme_list == []:
            raise ValueError("SummaryState appears to be empty. Please run the summarization process before instantiating the Render class.")
        
        if render_object is not None:
            if not isinstance(render_object, tuple) or len(render_object) != 2:
                raise ValueError("render_object must be a tuple of length 2.")
            if render_object[0] not in ["cluster_summary", "populated_theme", "redundancy"]:
                raise ValueError(f"Invalid render_object '{render_object[0]}'. Valid options are 'cluster_summary', 'populated_theme', or 'redundancy'.")

        self.summary_state = summary_state
        self.corpus_state = corpus_state
        self.render_object = render_object
        self.force = force
        self.llm_client = llm_client
        self.ai_model = ai_model
        self.render_hash = render_hash
        self.render_path = render_path
        self.output_save_location: str = output_save_location

        # Get the exact content that we want to render on, depending on whether the use ran the redundancy pass or stopped at the final orphan pass ()
        self.summary_to_render = self._get_summaries_for_render(force=force)
        # Add question_text to the summary_to_render so that it can be used as part of the payload for the question summaries and executive summary generation, and also added to the render_df for streaming out as part of the final document generation. We do this here so that the summary_to_render df is the complete df we need to work with for all subsequent operations, rather than having to merge in question text at multiple subsequent steps. This is a left merge because there may be some summaries that do not have question text if there are orphans that have not been handled, but we want to keep those summaries in the render_df and just have null question text for them.
        self.summary_to_render = (
            self.summary_to_render
            .merge(self.corpus_state.questions[["question_id", "question_text"]], 
                   on="question_id", 
                   how="left")
        )
        # Ensure sort for subsequent operations
        self.summary_to_render = self.summary_to_render.sort_values(by=["question_id", "theme_id"]).reset_index(drop=True)
        # Hash the summary to eiher check against previous renders or to set if this is the first time. This guards against changes in summary objects during rendering pass
        current_summary_hash = self._compute_df_hash(self.summary_to_render)

        # Handle check and recovery of pervious renders
        self._reinitialize_render(current_summary_hash=current_summary_hash)

    def _get_summaries_for_render(self, force) -> pd.DataFrame:
        """
        Determine which summaries to render based on the current stage of the summarization process, as indicated by summary_state.status(). If force is True, will render the summaries from the current stage even if the summarization process has not been completed. If render_object is specified, will render the specified object at the specified stage regardless of the current stage or force setting.
        """
        # First check whether the user has specified a render object, if so make sure they have set flag to true
        if self.render_object is not None and force == False:
            raise ValueError("You have specified a render_object but force is set to False. Please set force to True to render the specified object, or set render_object to None to automatically determine which summaries to render based on the current stage of the summarization process.")
        # If render object and force flag are set return the specified object at the specified stage, make sure index is within range
        if self.render_object is not None and force == True:
            obj_name, index = self.render_object
            if obj_name == "cluster_summary":
                if not (-len(self.summary_state.cluster_summary_list) <= index < len(self.summary_state.cluster_summary_list)):
                    raise ValueError(f"Index {index} out of range for cluster_summary_list with length {len(self.summary_state.cluster_summary_list)}.")
                return self.summary_state.cluster_summary_list[index]
            if obj_name == "populated_theme":
                if not (-len(self.summary_state.populated_theme_list) <= index < len(self.summary_state.populated_theme_list)):
                    raise ValueError(f"Index {index} out of range for populated_theme_list with length {len(self.summary_state.populated_theme_list)}.")
                return self.summary_state.populated_theme_list[index]
            else:
                obj_name = "redundancy" # This is not needed, just here for explicit clarity
                if not (-len(self.summary_state.redundancy_list) <= index < len(self.summary_state.redundancy_list)):
                    raise ValueError(f"Index {index} out of range for redundancy_list with length {len(self.summary_state.redundancy_list)}.")
                return self.summary_state.redundancy_list[index]


        # If render object not set, automatically get the summary object. First get the standard approach - redundancy or orphan - if they are complete. Otherwise check the flaf force option and return other options
        status_diagnostic = self.summary_state.status(diagnostic=True)
        max_stage, last_stage = status_diagnostic["max_stage"], status_diagnostic["stage"]
        # If the last stage is mapped themes or theme schema, we have to check whether a populated theme was generated - i.e. is max stage > 1 (i.e. a full pass has happened). If so the last available render is on the last populated_theme
        if last_stage in ["theme_schema", "mapped_theme"] and max_stage > 1:
            last_stage = "populated_theme"
        elif last_stage in ["theme_schema", "mapped_theme", "cluster_summary"] and max_stage <= 1:
            last_stage = "cluster_summary"
        else:
            last_stage = last_stage
        
        
        if last_stage == "redundancy":
            print("Summaries have gone through the redundancy pass and been checked for orphans. This is a readability first output, but has a risk of insight ommission. If fidelity is your priority you should run the orphan pass again before rendering. SummaryState.handle_orphans().")
            return(self.summary_state.redundancy_list[0])
        if last_stage == "orphan":
            print("Summaries have gone through the orphan pass, but not been checked for redundancy. This is a fidelity first output and may contain redundancies. If you want to handle redundancies before rendering, please run the redundancy pass: SummaryState.address_redundancy().")
            return(self.summary_state.populated_theme_list[-1])
        if last_stage not in ["redundancy", "orphan"] and not force:
            raise ValueError(f"Your summary state is currently at stage '{last_stage}'. You can only render summaries if you have completed the orphan pass and/or the redundancy pass. Use SummaryState.status() to determine your current stage. " 
                                "If you know what you are doing and want to render with the current summaries, please instantiate the Render class with force=True to override this check.")
        if last_stage in ["cluster_summary", "populated_theme"] and force:
            print(f"WARNING Force rendering with summaries from stage '{last_stage}'. Be aware that the summaries may not be fully processed and may contain redundancies or unhandled orphans.")
            if last_stage == "populated_theme":
                return(self.summary_state.populated_theme_list[-1])
            else:
                return(self.summary_state.cluster_summary_list[0])
            
    @staticmethod
    def _compute_df_hash(df) -> str:
        """
        Create a deterministic SHA-256 hash of a DataFrame.
        Used for state fingerprinting and resume validation.
        """

        if df is None or df.empty:
            return "EMPTY"

        # Defensive copy
        df_normalized = df.copy()

        # Normalize NaNs
        df_normalized = df_normalized.fillna("__NULL__")

        # Sort columns alphabetically for consistency
        df_normalized = df_normalized.sort_index(axis=1)

        # Sort rows deterministically by all columns
        df_normalized = df_normalized.sort_values(
            by=list(df_normalized.columns)
        ).reset_index(drop=True)

        # Convert to stable JSON representation
        json_bytes = df_normalized.to_json().encode("utf-8")

        # Hash
        return hashlib.sha256(json_bytes).hexdigest()          
                   
    def _reinitialize_render(self, current_summary_hash) -> None:
        """
        Checks if there is another render hash file saved. If so, checks the hashes to ensure match
        If conflict force abort or overwrite
        If hash match, reload the render df and any attributes created thus far
        If no previous render save the current hash and create the clean render_df from the summary_to_render   

        """
        # Check if any previous render objects exist:
        os.makedirs(self.render_path, exist_ok=True) # Ensure the render path exists

        if os.path.exists(os.path.join(self.render_path, self.render_hash)):
            saved_hash_df = pd.read_parquet(os.path.join(self.render_path, self.render_hash))
            saved_hash = saved_hash_df["summary_hash"].iloc[0]
            if saved_hash != current_summary_hash:
                restart = None
                while restart not in ["1", "2"]:
                    restart = input(
                        "A previous render hash was found that does not match the current summary hash. This suggests that the summaries have changed since the last render. " \
                        "Do you want to proceed with rendering and overwrite the previous render?\n"
                        "(1) Yes, proceed with rendering and overwrite the previous render.\n" 
                        "(2) No, stop the rendering process to avoid overwriting the previous render (this will abort render and require the correct summary_state be passed to Render init).\n"
                        "(1/2):\n"
                        ).strip()

                if restart == "2":
                    raise ValueError("Rendering process aborted by user to avoid overwriting previous render. Please check your summary_state and ensure it is correct before re-instantiating the Render class.")
                else:
                    # Clean up the render artifacts and start new render
                    os.remove(os.path.join(self.render_path, self.render_hash))
                    if os.path.exists(os.path.join(self.render_path, self.render_df)):
                        os.remove(os.path.join(self.render_path, self.render_df)) 
                    # Create and save the new hash file
                    new_hash_df = pd.DataFrame({"summary_hash": [current_summary_hash]})
                    new_hash_df.to_parquet(os.path.join(self.render_path, self.render_hash), index=False)
                    print("Previous render artifacts deleted. Hash set to current summary_to_render.")
                    return(None)
                
            else: # if the hash matches we load the old dataframe and create all the attributes that we can from it.
                print("A previous render with the same summary hash was found. Loading the previous generate cosmetic artifacts\n")
                path = Path(self.render_path)
                loaded_artefacts = []
                for file in path.glob("*.parquet"):
                    if file.stem == config.render_prefix["render_df"]:
                        self.render_df = pd.read_parquet(file)
                        loaded_artefacts.append("render_df")
                    elif file.stem == config.render_prefix["render_title_exec_summary_df"]:
                        self.title_exec_summary_df = pd.read_parquet(file)
                        loaded_artefacts.append("title_summary_df")
                    elif file.stem == config.render_prefix["render_question_summary_df"]:
                        self.question_summary_df = pd.read_parquet(file)
                        loaded_artefacts.append("question_summary_df")
                    else:
                        if file.stem == config.render_prefix["render_sylized_rewrite_df"]:
                            self.stylized_rewrite_df = pd.read_parquet(file)
                            loaded_artefacts.append("stylized_rewrite_df")
                
                if not hasattr(self, "render_df"):
                    print(f"Cosmetic artifacts loaded: {loaded_artefacts}, but final render_df not found. You should add any missing artifacts and call integrat_cosmetic_artifacts() to construct the render_df)")


        else: # if no previous file exists, save the hash and create the render df from the summary_to_render
            # Create and save the new hash file
            new_hash_df = pd.DataFrame({"summary_hash": [current_summary_hash]})
            new_hash_df.to_parquet(os.path.join(self.render_path, self.render_hash), index=False)
            return(None)

    def _gen_question_payload(self, qid) -> str:
        """
        Generates a string that can be used as the payload to the llm call.
        The string is a concat of the themmatic summaries preceeded with the question text and with clear boundaries between the themes
        It can be used as the payload for the question summries and as part of the exec summaries
        """
        question_df = self.summary_to_render[self.summary_to_render["question_id"] == qid].copy()
        output = ""
        for _, row in question_df.iterrows():
            output += f"Theme: {row['theme_label']}\n"
            output += f"{row['thematic_summary']}\n"
            output += "--- END THEME ---\n"

        return(output)
    
    
    def gen_question_summaries(self) -> pd.DataFrame:

        # First check whether question summaries already exist, if so offer reload or regenerate options
        if hasattr(self, "question_summary_df"):
            rerun = None
            while rerun not in ["1", "2"]:
                rerun = input(
                    "Question summaries already exist. Do you want to regenerate the question summaries?\n"
                    "(1) Yes, regenerate the question summaries.\n" 
                    "(2) No, keep the existing question summaries.\n"
                    "(1/2):\n"
                ).strip()
            if rerun == "2":
                print("Keeping existing question summaries.")
                return self.question_summary_df
            else:                
                print("Regenerating question summaries.")

        # Then check whether a stylistic re-write has been done. Flag for the user that if they want a stylistic re-write its likely better to do it after the summaries have been generated, but they can proceed as they like
        if not hasattr(self, "stylized_rewrite_df"):
            abort = None
            while abort not in ["1", "2"]:
                abort = input(
                    "You have not done a stylistic rewrite of the themes. This is not neccesary to proceed, "
                    "however if you intend to do a stylistic rewrite and want  the style of the re-write to appear in your summaries it is reccomended to do the re-write first. "
                    "Do you want to proceed with generating question summaries without doing a stylistic rewrite?\n"
                    "(1) Yes, proceed with generating question summaries without doing a stylistic rewrite.\n" 
                    "(2) No, I want to do a stylistic rewrite before generating question summaries.\n"
                    "(1/2):\n"
                ).strip()

                if abort == "2":
                    print("Aborting question summary generation so that you can do a stylistic rewrite first. Please run stylistic_rewrite() and then re-run gen_question_summaries() to generate the question summaries with the rewritten themes.")
                    return(self.render_df)
        
        # loop over question id to get the question payload
        question_summaries_df_list = []

        working_render_df = self.summary_to_render.copy()

        for qid in working_render_df["question_id"].unique():
            question_id = qid
            question_text = working_render_df[working_render_df["question_id"] == qid]["question_text"].iloc[0]
            question_payload = self._gen_question_payload(qid=question_id)

        # Prepare the full payload for the llm call
            sys_prompt = Prompts().question_summaries()
            user_prompt = (
            f"Research question: {question_text}\n\n"
            "CONTENT TO SUMMARIZE:\n"
            f"{question_payload}\n\n"
            )

            fall_back = {"summary": ""}
            json_schema = {
                "name": "question_summary_generator",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "summary": {
                            "type": "string",
                            "description": "A single-paragraph thematic overview (3-5 sentences) synthesizing the themes for the research question."
                        }
                    },
                    "required": ["summary"],
                    "additionalProperties": False
                }
            }
            
            # Call the llm
            response = utils.call_chat_completion(
                sys_prompt=sys_prompt,
                user_prompt=user_prompt,
                llm_client=self.llm_client,
                ai_model=self.ai_model,
                return_json=True,
                json_schema=json_schema,
                fall_back=fall_back
            )

            question_summary_txt = response["summary"]
            question_summary_df = pd.DataFrame({
                "content": [question_summary_txt],
                "doc_attr": ["question_summary"],
                "question_id": [qid],
            })

            question_summaries_df_list.append(question_summary_df)

        question_summaries_df_out = pd.concat(question_summaries_df_list, ignore_index=True).sort_values(by=["question_id"]).reset_index(drop=True)

        self.question_summary_df = question_summaries_df_out
        self.question_summary_df.to_parquet(os.path.join(self.render_path, config.render_prefix["render_question_summary_df"] + ".parquet"), index=False)       
        return(self.question_summary_df)
    
    
    def _gen_exec_summary_payload(self) -> str:
        temp_summary_df = self.summary_to_render.copy()
        temp_summary_df = temp_summary_df.sort_values(by=["question_id", "theme_id"]).reset_index(drop=True)

        # Loop over the questions
        output = ""
        for qid in temp_summary_df["question_id"].unique():
            output += f"Research Question: {temp_summary_df[temp_summary_df['question_id'] == qid]['question_text'].iloc[0]}\n"
            question_output = self._gen_question_payload(qid)
            output += question_output
            output += "\n=== END QUESTION ===\n\n"

        return output
    

    def gen_exec_summary(self, word_length: int = 500) -> str:

        if hasattr(self, "title_exec_summary_df"):
            rerun = None
            while rerun not in ["1", "2"]:
                rerun = input(
                    "An executive summary and title have already been generated. Do you want to regenerate the executive summary and title?\n"
                    "(1) Yes, regenerate the executive summary and title.\n" 
                    "(2) No, keep the existing executive summary and title.\n"
                    "(1/2):\n"
                ).strip()
                
            if rerun == "2":
                print("Keeping existing executive summary and title.")
                return self.title_exec_summary_df
            else:                
                print("Regenerating executive summary and title.")

        exec_summary_payload = self._gen_exec_summary_payload()
        
        sys_prompt = Prompts().exec_summary(word_length=word_length)
        user_prompt = exec_summary_payload
        fall_back = {"executive_summary": "", "title": ""}
        json_schema = {
            "name": "executive_summary_generator",
            "strict": True,
            "schema": {
                "type": "object",
                "properties": {
                    "executive_summary": {
                        "type": "string",
                        "description": "The final executive summary text in continuous prose."
                    },
                    "title": {
                        "type": "string",
                        "description": "A concise descriptive title, maximum 12 words, no subtitle."
                    }
                },
                "required": ["executive_summary", "title"],
                "additionalProperties": False
            }
        }

        response = utils.call_chat_completion(
            sys_prompt=sys_prompt,
            user_prompt=user_prompt,
            llm_client=self.llm_client,
            ai_model=self.ai_model,
            return_json=True,
            json_schema=json_schema,
            fall_back=fall_back
        )

        exec_summary = response["executive_summary"]
        title = response["title"]
        
        title_exec_summary_df = pd.DataFrame({
            "content": [title, exec_summary],
            "doc_attr": ["title", "exec_summary"]
            })

        self.title_exec_summary_df = title_exec_summary_df

        self.title_exec_summary_df.to_parquet(os.path.join(self.render_path, config.render_prefix["render_title_exec_summary_df"] + ".parquet"), index=False)

        return self.title_exec_summary_df

    
    def stylistic_rewrite(self, style: str = "academic") -> pd.DataFrame:
        """
        Rewrites the themes in render_df in a specific style. Default is academic, but this can be varied (e.g. journalistic, conversational, engaging,etc) depending on user preference and intended audience. The rewrite will attempt to maintain the original meaning and insights while improving readability and engagement.
        Also goes about varying the language to make the content more engaging and less monotonous. 
        While this re-write intends to maintain complete fidelity any re-write risks some insight loss. 
        This is an optional pass for the user if they want to prioritize readability over fidelity. 
        """
        # First check wheter a stylistic rewrite has already been done, if so offer reload or regenerate options
        if hasattr(self, "stylized_rewrite_df"): 
            rerun = None
            while rerun not in ["1", "2"]:
                rerun = input(
                    "A stylistic rewrite of the themes has already been generated. Do you want to regenerate the stylistic rewrite?\n"
                    "(1) Yes, regenerate the stylistic rewrite.\n" 
                    "(2) No, keep the existing stylistic rewrite.\n"
                    "(1/2):\n"
                ).strip()
                
            if rerun == "2":
                print("Keeping existing stylistic rewrite.")
                return self.stylized_rewrite
            else:                
                print("Regenerating stylistic rewrite.")
        
        # Then check whether the summaries have already been generated. If they have, send a warning that it is better to run the stylistic re-write prior to summarizing so that the tone is carried through
        if hasattr(self, "question_summary_df"):
            abort = None
            while abort not in ["1", "2"]:
                abort = input(
                    "You have already generated question summaries. It is generally recommended to do the stylistic rewrite before generating the question summaries so that the style is carried through to the question summaries. Do you want to proceed with the stylistic rewrite even though you have already generated question summaries?\n"
                    "(1) Yes, proceed with the stylistic rewrite even though I have already generated question summaries.\n" 
                    "(2) No, I want to do a stylistic rewrite before generating question summaries.\n"
                    "(1/2):\n"
                ).strip()

                if abort == "2":
                    print("Aborting stylistic rewrite so that you can do it before generating question summaries. Please run stylistic_rewrite() and then re-run gen_question_summaries() to generate the question summaries with the rewritten themes.")
                    return(self.render_df)
        
        stylized_content_df_out = pd.DataFrame(columns=["stylized_text", "question_id", "theme_id"])

        for qid in self.summary_to_render["question_id"].unique():
            question_df = self.summary_to_render[self.summary_to_render["question_id"] == qid].copy()
            # Place frozen content here so that it resets for each rq
            frozen_content = ""
            for idx, row in question_df.iterrows():
                question_text = row["question_text"]
                content = row["thematic_summary"]
                theme_label = row["theme_label"]
                theme_id = row["theme_id"]
                sys_prompt = Prompts().stylistic_rewrite(style=style, index = idx)
                user_prompt = (
                    f"CURRENT RESEARCH QUESTION: {question_text}\n"
                    "FROZEN CONTENT:\n"
                    f"{frozen_content}\n\n"
                    f"CURRENT THEME LABEL: {theme_label}\n"
                    "THEMATIC SUMMARY TO STYLE:\n"
                    f"{content}\n\n"
                )
                fall_back = {"refined_summary": ""}

                json_schema = {
                    "name": "stylistic_rewrite_generator",
                    "strict": True,
                    "schema": {
                        "type": "object",
                        "properties": {
                            "refined_summary": {
                                "type": "string",
                                "description": "Refined thematic summary text."
                            }
                        },
                        "required": ["refined_summary"],
                        "additionalProperties": False
                    }
                }
                response = utils.call_chat_completion(
                    sys_prompt=sys_prompt,
                    user_prompt=user_prompt,
                    llm_client=self.llm_client,
                    ai_model=self.ai_model,
                    return_json=True,
                    json_schema=json_schema,
                    fall_back=fall_back
                )

                # Get the data, transform to df
                stylized_content = response["refined_summary"]
                stylized_content_df = pd.DataFrame({
                    "stylized_text": [stylized_content],
                    "question_d_id": [qid],
                    "theme_id": [theme_id],
                    })
                
                # Append to the output df
                stylized_content_df_out = pd.concat([stylized_content_df_out, stylized_content_df], ignore_index=True)
                frozen_content += f"{row['theme_label']}\n\n{stylized_content}\n\n"
        
        self.stylized_rewrite_df = stylized_content_df_out
        self.stylized_rewrite_df.to_parquet(os.path.join(self.render_path, config.render_prefix["render_sylized_rewrite_df"] + ".parquet"), index=False)
        return(stylized_content_df_out)
    
    def _add_question_summaries_to_render_df(self, render_df, question_summaries: list) -> pd.DataFrame:
        """
        Add question summaries to the render dataframe.
        Outputs a dataframe with questions and summaries
        This will still have to be integrated title and exec summary from the in render_df - this check happens in gen_question_summaries
        """
        working_render_df = render_df.copy()

        # Separate the working_render_df into a list of dataframes, one for each question
        list_of_qdfs = []
        for qid in self.corpus_state.questions["question_id"].to_list():
            question_df = working_render_df[working_render_df["question_id"] == qid].copy()
            list_of_qdfs.append(question_df)
        
        # Concat the question summaries with the corresponding question dfs and create a new list of dataframes with summaries
        list_of_qdfs_with_summaries = []
        for qdf, q_summary in zip(list_of_qdfs, question_summaries):
            question_summary_df = pd.DataFrame({
                "content": [q_summary],
                "doc_attr": ["question_summary"],
                "question_id": [qdf["question_id"].iloc[0]]
            })
            qdf_with_summary = pd.concat([question_summary_df, qdf], ignore_index=True)
            list_of_qdfs_with_summaries.append(qdf_with_summary)

        # Concat the list of question dfs with summaries into a single dataframe
        questions_with_summaries_df = pd.concat(list_of_qdfs_with_summaries, ignore_index=True)
        # Return the resulting dataframe
        return questions_with_summaries_df


    def integrate_cosmetic_changes(self) -> pd.DataFrame:

        # --- Hash validation ---
        current_summary_hash = self._compute_df_hash(self.summary_to_render)
        saved_hash_df = pd.read_parquet(os.path.join(self.render_path, self.render_hash))
        saved_hash = saved_hash_df["summary_hash"].iloc[0]

        if current_summary_hash != saved_hash:
            raise ValueError(
                "The summary_to_render has changed since Render initialization. "
                "Reinitialize Render before integrating artifacts."
            )

        # --- Base render df ---
        render_df = (
            self.summary_to_render
            .sort_values(by=["question_id", "theme_id"])
            .copy()
        )

        render_df.rename(columns={"thematic_summary": "content"}, inplace=True)
        render_df["doc_attr"] = "thematic_summary"

        # --- Question summaries ---
        if hasattr(self, "question_summary_df"):
            render_df = self._add_question_summaries_to_render_df(
                render_df,
                self.question_summary_df
            )

        # --- Stylized rewrite ---
        if hasattr(self, "stylized_rewrite_df"):
            render_df = render_df.merge(
                self.stylized_rewrite_df[["question_id", "theme_id", "stylized_text"]],
                on=["question_id", "theme_id"],
                how="left"
            )
            render_df["stylized_text"] = render_df["stylized_text"].fillna(render_df["content"])

        # --- Title + Exec summary ---
        if hasattr(self, "title_exec_summary_df"):
            title_exec_df = self.title_exec_summary_df.copy()
            title_exec_df["question_id"] = None
            title_exec_df["theme_id"] = None
            render_df = pd.concat([title_exec_df, render_df], ignore_index=True)

        # --- Final ordering ---
        render_df["doc_order"] = range(len(render_df))

        # --- Save ---
        self.render_df = render_df
        render_df.to_parquet(
            os.path.join(self.render_path, config.render_prefix["render_df"] + ".parquet"),
            index=False
        )
        return self.render_df
    
    def render_output(
        self,
        output_type: str = "md",
        use_stylized: bool = False,
        filename: str = "rendered_output"
    ) -> None:
        """
        Renders the integrated render_df into the selected output format.

        Args:
            output_type: "md", "docx", or "pdf"
            use_stylized: If True uses stylized_text when available.
            filename: Base filename (no extension)
        """

        if not hasattr(self, "render_df"):
            raise ValueError(
                "render_df not found. Please run integrate_cosmetic_changes() first."
            )

        if output_type not in ["md", "docx", "pdf"]:
            raise ValueError("output_type must be one of: 'md', 'docx', 'pdf'.")

        df = self.render_df.copy()

        # Choose text source
        if use_stylized and "stylized_text" in df.columns:
            df["final_text"] = df["stylized_text"].fillna(df["content"])
        else:
            df["final_text"] = df["content"]

        # Dispatch
        if output_type == "md":
            self._render_to_markdown(df, filename)

        elif output_type == "docx":
            self._render_to_docx(df, filename)

        elif output_type == "pdf":
            self._render_to_pdf(df, filename)

        print(f"Render complete: {filename}.{output_type}")


    def _render_to_markdown(self, df, filename):

        lines = []

        for _, row in df.sort_values("doc_order").iterrows():

            text = row["final_text"]
            attr = row["doc_attr"]

            if attr == "title":
                lines.append(f"# {text}\n")

            elif attr == "exec_summary":
                lines.append(f"{text}\n")

            elif attr == "question_summary":
                lines.append(f"## Question Overview\n{text}\n")

            elif attr == "thematic_summary":
                lines.append(f"### {row.get('theme_label','')}\n{text}\n")

        output_path = os.path.join(self.output_save_location, f"{filename}.md")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _render_to_docx(self, df, filename):

        document = Document()

        for _, row in df.sort_values("doc_order").iterrows():

            text = row["final_text"]
            attr = row["doc_attr"]

            if attr == "title":
                document.add_heading(text, level=0)

            elif attr == "exec_summary":
                document.add_paragraph(text)

            elif attr == "question_summary":
                document.add_heading("Question Overview", level=1)
                document.add_paragraph(text)

            elif attr == "thematic_summary":
                document.add_heading(row.get("theme_label",""), level=2)
                document.add_paragraph(text)

        output_path = os.path.join(self.output_save_location, f"{filename}.docx")
        document.save(output_path)

    def _render_to_docx(self, df, filename):

        document = Document()

        for _, row in df.sort_values("doc_order").iterrows():

            text = row["final_text"]
            attr = row["doc_attr"]

            if attr == "title":
                document.add_heading(text, level=0)

            elif attr == "exec_summary":
                document.add_paragraph(text)

            elif attr == "question_summary":
                document.add_heading("Question Overview", level=1)
                document.add_paragraph(text)

            elif attr == "thematic_summary":
                document.add_heading(row.get("theme_label",""), level=2)
                document.add_paragraph(text)

        output_path = os.path.join(self.output_save_location, f"{filename}.docx")
        document.save(output_path)

    
    
    
    
    
    
    
    # def summary_to_doc(self, paper_title: str = None, summary_filename: str = None) -> str:
        
    #     def _sanitize(text: str) -> str:
    #         control_chars = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F]')
    #         if text is None:
    #             return ""
    #         # ensure str
    #         s = str(text)
    #         # normalize unicode
    #         s = unicodedata.normalize("NFC", s)
    #         # drop XML-illegal control chars (keep \t, \n, \r)
    #         s = control_chars.sub('', s)
    #         return s
        
    #     if summary_filename is None:
    #         summary_filename = "literature_review.docx"

    #     if paper_title is None and hasattr(self, "title"):
    #         paper_title = self.title
    #     elif paper_title is None:
    #         raise ValueError("Paper_title must be provided. Either pass it in summary_to_doc or call get_executive_summary() which will provide a title.")

    #     doc = Document()
    #     doc.add_heading(_sanitize(paper_title), level=0)

    #     # Executive summary
    #     if getattr(self, "exec_summary", None):
    #         doc.add_heading("Executive summary", level=1)
    #         for para in str(self.exec_summary).splitlines():
    #             p = _sanitize(para)
    #             if p.strip():
    #                 doc.add_paragraph(p)
    #         doc.add_page_break()
    #     else: 
    #         raise ValueError("exec_summary attribute not found. Please run gen_executive_summary() before exporting to doc.")

    #     df = self.summaries.copy().reset_index(drop=False).rename(columns={"index": "_row"})
    #     themed = {"label", "contents"}.issubset(df.columns)

    #     if themed:
    #         for question_text, qdf in df.groupby("question_text", sort=False):
    #             doc.add_heading(_sanitize(question_text), level=1)
    #             if "question_summary" in qdf.columns:
    #                 qs = (qdf["question_summary"].dropna().astype(str).iloc[0]
    #                     if not qdf["question_summary"].dropna().empty else "")
    #                 if qs.strip():
    #                     doc.add_paragraph(_sanitize(qs))
    #             qdf = qdf.sort_values("_row")
    #             for _, r in qdf.iterrows():
    #                 label = _sanitize(r.get("label", ""))
    #                 contents = _sanitize(r.get("contents", ""))
    #                 if label:
    #                     doc.add_heading(f"Theme: {label}", level=2)
    #                 if contents:
    #                     for para in contents.splitlines():
    #                         p = _sanitize(para)
    #                         if p.strip():
    #                             doc.add_paragraph(p)
    #             doc.add_page_break()
    #     else:
    #         for question_text, qdf in df.groupby("question_text", sort=False):
    #             doc.add_heading(_sanitize(question_text), level=1)
    #             if "question_summary" in qdf.columns:
    #                 qs = (qdf["question_summary"].dropna().astype(str).iloc[0]
    #                     if not qdf["question_summary"].dropna().empty else "")
    #                 if qs.strip():
    #                     doc.add_paragraph(_sanitize(qs))
    #             qdf = qdf.sort_values("_row")
    #             for _, r in qdf.iterrows():
    #                 summ = _sanitize(r.get("summary", ""))
    #                 if summ:
    #                     for para in summ.splitlines():
    #                         p = _sanitize(para)
    #                         if p.strip():
    #                             doc.add_paragraph(p)
    #             doc.add_page_break()

        
    #     # Save
    #     save_dir = self.output_save_location
    #     os.makedirs(save_dir, exist_ok=True)
        
    #     out_path = os.path.join(save_dir, summary_filename)
    #     doc.save(out_path)
        
    #     # Mark that doc has been generated - as flag for running the AI peer review (neccesary as all the steps for running the dog generation are needed for peer reviewn).
    #     self.doc = True

    #     return (
    #         f'Word doc of the literature review generated and save here: {out_path}'
    #     )
    
    # def get_ai_peer_review(self, 
    #                        save_directory: str = None,
    #                        save_filename: str = "ai_peer_review.parquet",
    #                        output_length: int = 5000, 
    #                        max_tokens: int = 10000) -> pd.DataFrame:
    #     """
    #     Request an AI peer review of the concatenated summaries.

    #     Args:
    #         output_length: Suggested maximum word count for the review.
    #         max_tokens: Hard limit on token usage for the AI model.

    #     Returns:
    #         AI review as a dataframe. If the review exceeds token budget, 'error' key may appear.
    #     """        
    #     # Function specific utilities ---------------------
    #     # Get the RQ and summaries excluding the current RQ
    #     def get_rq_summaries(self, current_rq_id) -> str:
    #         output_parts = []
    #         for rq_id, rq_df in self.summaries.groupby("question_id", sort=False):
    #             if rq_id == current_rq_id:
    #                 continue
    #             else:
    #                 output_parts.append(f"RESEARCH QUESTION:\n{rq_id}\n")
    #                 question_summary = rq_df["question_summary"].iloc[0]
    #                 output_parts.append(f"SUMMARY:\n{question_summary}\n")
    #         return "\n".join(output_parts)
        
    #     # Get the RQ and its associated content
    #     def get_rq_content(self, current_rq_id) -> str:
    #         current_rq_df = self.summaries[self.summaries["question_id"] == current_rq_id]
    #         current_rq_text = current_rq_df["question_text"].iloc[0]
    #         output_parts = []
    #         output_parts.append(f"RESEARCH QUESTION:\n{current_rq_text}\n")
    #         themed = {"label", "contents"}.issubset(self.summaries.columns)
    #         if themed:
    #             for _, row in current_rq_df.iterrows():
    #                 label = row["label"].strip()
    #                 content = row["contents"].strip()
    #                 output_parts.append(
    #                     f"Theme: {label}\n"
    #                     f"Summary: {content}\n"
    #                 )

    #             return "\n".join(output_parts)
            
    #         else:
    #             for _, row in current_rq_df.iterrows():
    #                 summary = row["summary"].strip()
    #                 output_parts.append(f"Summary: {summary}\n")
    #             return "\n".join(output_parts)
        
    #     # END UTILS -------------------------------

    #     # Check that executive summary has been generated
    #     if getattr(self, "exec_summary", None) is None:
    #         raise ValueError("Please run gen_executive_summary() before requesting an AI peer review.")
        
    #     # Check that the peer review has not already been generated
    #     if save_directory is None:
    #         save_directory = self.summaries_folder

    #     if os.path.exists(os.path.join(save_directory, save_filename)):
    #         recover = None
    #         while recover not in ["r", "n"]:
    #             recover = input("AI peer review file already exists. Loading existing file. (r)eload or create (n)ew? ")
    #         if recover == "r":
    #             self.ai_peer_review = pd.read_parquet(os.path.join(save_directory, save_filename))
    #             return self.ai_peer_review
    #         else:
    #             print("Generating new AI peer review and overwriting existing file.")

    #     # Populate the list that will form the prompt for the LLM
    #     output_list = []
    #     for rq_id, rq_df in self.summaries.groupby("question_id", sort=False):
    #         output_parts = []
    #         output_parts.append("INFORMATION FOR CONTEXT\n")
    #         output_parts.append("EXECUTIVE SUMMARY:\n")
    #         output_parts.append(self.exec_summary + "\n")
    #         output_parts.append("OTHER RESEARCH QUESTIONS AND SUMMARIES:\n")
    #         output_parts.append(get_rq_summaries(self, current_rq_id=rq_id))
    #         output_parts.append("CURRENT RESEARCH QUESTION AND CONTENT:\n")
    #         output_parts.append(get_rq_content(self, current_rq_id=rq_id))

    #         output_string = "\n".join(output_parts)
            
    #         # Call the reasoning model
    #         resp = utils.call_reasoning_model(
    #             prompt=Prompts().ai_peer_review(
    #                 lit_review=output_string,
    #                 output_length=output_length,
    #                 max_tokens=max_tokens,
    #                 themed={"label", "contents"}.issubset(self.summaries.columns)
    #             ),
    #             llm_client=self.llm_client,
    #             ai_model=self.ai_model,
    #             timeout=1200
    #         )

    #         # Handle the response with a retry to clean the JSON via an LLM
    #         try:
    #             response_dict = json.loads(resp)
    #         except json.JSONDecodeError:
    #             response_retry = utils.llm_json_clean(x = resp,
    #                                             prompt = Prompts().peer_review_format_check(),
    #                                             llm_client=self.llm_client,
    #                                             ai_model=self.ai_model)
    #             response_dict = json.loads(response_retry)

    #         overall_comment_df = pd.DataFrame({"comment_id": "overall", 
    #                                            "comment": response_dict.get("overall_comment"),
    #                                            "severity": pd.NA, 
    #                                            "location": "overall"}, index=[0])
                 

    #         specific_comments_df = pd.DataFrame(
    #             response_dict.get("specific_comments", []), 
    #             columns=["comment_id", "comment", "severity", "location"]
    #         )


    #         output = pd.concat([overall_comment_df, specific_comments_df], ignore_index=True)
    #         output["research_question_id"] = rq_id
    #         output["research_question_text"] = rq_df["question_text"].iloc[0]
    #         output["resubmit"] = response_dict.get("resubmit", False)

    #         output_list.append(output)

    #     output_df = pd.concat(output_list).reset_index(drop=True)
    #     # Convert to string and fill missing severity with "None" so that parquet will accept it
    #     output_df["severity"] = output_df["severity"].apply(lambda x: x if pd.notna(x) else "None")
    #     # Also make all id str (handle int and "overall" fields)
    #     output_df["comment_id"] = output_df["comment_id"].astype(str)

    #     self.ai_peer_review = output_df

    #     os.makedirs(save_directory, exist_ok=True)
    #     self.ai_peer_review.to_parquet(os.path.join(save_directory, save_filename), index=False)
        
    #     return self.ai_peer_review
    
    # def peer_review_to_doc(self):
    #     if self.ai_peer_review is None:
    #         print("No AI peer review data available. Please run get_ai_peer_review() first.")
    #         return None
        
    #     # Convert the AI peer review DataFrame to a formatted string or document
    #     # This is a placeholder for actual document generation logic
    #     doc = Document()
    #     for _, row in self.ai_peer_review.iterrows():
    #         doc.add_paragraph(f"Comment ID: {row['comment_id']}")
    #         if row["severity"] is not None:
    #             doc.add_paragraph(f"Severity: {row['severity']}")
    #         if row["resubmit"]:
    #             doc.add_paragraph("**Resubmission Required**")
    #         if row["location"] == "overall":
    #             doc.add_paragraph(f"Location: {row['research_question_text']}")
    #         else:
    #             doc.add_paragraph(f"Location: {row['research_question_text']} - {row['location']}")
    #         doc.add_paragraph(f"Comment: {row['comment']}")
    #         doc.add_paragraph("---------------------")


    #     # Save
    #     save_dir = self.output_save_location
    #     os.makedirs(save_dir, exist_ok=True)

    #     out_path = os.path.join(save_dir, "peer_review.docx")
    #     doc.save(out_path)

    #     print(f"Peer review generated and saved to {out_path}")
        